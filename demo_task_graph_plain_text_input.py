#!/usr/bin/env python3
"""
Task-graph ontology-guided PDF/DOCX/TXT/PLAIN-TEXT -> slide deck pipeline
======================================================

This version deliberately avoids generating the whole slide deck in one shot.
It decomposes the work into explicit tasks, completes each task independently,
validates/repairs each task result, and only then aggregates slide fragments into
one PML deck.

Core idea
---------
    PDF/DOCX/TXT report text
      -> report ontology guides section-level summarization
      -> slide ontology + layout registry guide task creation
      -> each task creates exactly one slide-spec / PML fragment
      -> each fragment is validated/repaired independently
      -> final aggregation builds the deck-level PML
      -> existing renderer produces HTML/PPTX

Install
-------
    pip install -U langgraph google-genai jinja2 python-pptx python-docx pypdf

Mock run
--------
    python demo_task_graph_docx_input.py dummy.docx \
      --renderer demo_dsl_conclusion_box.py \
      --out-dir out_task_graph_demo \
      --mock

Gemini run
----------
    export GEMINI_API_KEY="your-key"
    python demo_task_graph_docx_input.py input.docx \
      --renderer demo_dsl_conclusion_box.py \
      --out-dir out_task_graph_demo

Outputs
-------
    source_text.txt
    generated.pml
    corporate.psl
    task_plan.json
    task_outputs/*.json
    validation_report.json
    repair_log.txt
    output.html
    output.pptx
"""

from __future__ import annotations

import argparse
import importlib.util
import json
import os
import re
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional, TypedDict

from jinja2 import Environment, StrictUndefined


def require_langgraph():
    try:
        from langgraph.graph import END, START, StateGraph
    except ImportError as exc:
        raise RuntimeError("Missing dependency: langgraph. Install with: pip install -U langgraph") from exc
    return END, START, StateGraph


class PipelineState(TypedDict, total=False):
    pdf_path: str
    input_text: str
    use_stdin: bool
    renderer_path: str
    out_dir: str
    model: str
    mock: bool

    report_ontology_path: str
    slide_ontology_path: str
    layout_registry_path: str

    report_ontology: Dict[str, Any]
    slide_ontology: Dict[str, Any]
    layout_registry: Dict[str, Any]

    source_text: str
    report_summary: Dict[str, Any]

    task_plan: Dict[str, Any]
    pending_tasks: List[Dict[str, Any]]
    current_task: Dict[str, Any]
    current_task_output: Dict[str, Any]
    completed_tasks: List[Dict[str, Any]]

    pml_text: str
    psl_text: str

    pml_path: str
    psl_path: str
    task_plan_path: str
    html_path: str
    pptx_path: str
    md_path: str
    source_text_path: str

    validation_reports: List[Dict[str, Any]]
    repair_log: List[str]


DEFAULT_REPORT_ONTOLOGY: Dict[str, Any] = {
    "ontology_id": "report-domain-ontology-v0.2",
    "description": "Ontology nghiệp vụ báo cáo để chia nội dung thành các đề mục ổn định trước khi tạo slide.",
    "concepts": [
        {
            "id": "executive_summary",
            "preferred_heading": "Tóm tắt điều hành",
            "intent": "summarize_core_message",
            "max_bullets": 4,
            "slide_priority": 1,
        },
        {
            "id": "context",
            "preferred_heading": "Bối cảnh và vấn đề",
            "intent": "explain_context",
            "max_bullets": 5,
            "slide_priority": 2,
        },
        {
            "id": "key_findings",
            "preferred_heading": "Phát hiện chính",
            "intent": "present_findings",
            "max_bullets": 6,
            "slide_priority": 3,
        },
        {
            "id": "risks",
            "preferred_heading": "Rủi ro và điểm cần chú ý",
            "intent": "highlight_risks",
            "max_bullets": 5,
            "slide_priority": 4,
        },
        {
            "id": "recommendations",
            "preferred_heading": "Khuyến nghị hành động",
            "intent": "recommend_actions",
            "max_bullets": 5,
            "slide_priority": 5,
        },
        {
            "id": "roadmap",
            "preferred_heading": "Lộ trình triển khai",
            "intent": "show_progression",
            "max_bullets": 5,
            "slide_priority": 6,
        },
    ],
}


DEFAULT_SLIDE_ONTOLOGY: Dict[str, Any] = {
    "ontology_id": "slide-creation-ontology-v0.2",
    "concepts": {
        "SlideTask": {
            "description": "Một nhiệm vụ độc lập tạo một slide hoặc một fragment slide.",
            "required_fields": ["task_id", "source_concept_id", "slide_intent", "layout_id", "title"],
        },
        "SlideSpec": {
            "description": "Kết quả có cấu trúc sau khi hoàn thành một SlideTask.",
            "required_fields": ["title", "layout", "blocks"],
        },
    },
    "layout_selection_rules": [
        {"intent": "summarize_core_message", "prefer": ["title-bullets", "numbered-columns-3"]},
        {"intent": "explain_context", "prefer": ["title-bullets", "text-image"]},
        {"intent": "present_findings", "prefer": ["grid-4", "title-bullets", "numbered-columns-4"]},
        {"intent": "highlight_risks", "prefer": ["title-bullets", "grid-3"]},
        {"intent": "recommend_actions", "prefer": ["numbered-columns-4", "title-bullets"]},
        {"intent": "show_progression", "prefer": ["timeline", "stair-progress", "title-bullets"]},
    ],
    "task_policy": {
        "one_concept_one_task": True,
        "validate_each_task_before_aggregation": True,
        "aggregate_only_completed_tasks": True,
    },
}


DEFAULT_LAYOUT_REGISTRY: Dict[str, Any] = {
    "registry_id": "renderer-layout-registry-v0.2",
    "layouts": [
        {
            "id": "title-bullets",
            "intent": ["summarize_core_message", "explain_context", "highlight_risks"],
            "supported_blocks": ["title", "subtitle", "bullets", "conclusion", "footer_text"],
            "capacity": {"max_bullets": 7, "max_chars_per_bullet": 130},
            "description": "Tiêu đề + danh sách bullet; layout an toàn nhất cho nội dung báo cáo.",
        },
        {
            "id": "grid-4",
            "intent": ["present_findings"],
            "supported_blocks": ["title", "cells", "footer_text"],
            "capacity": {"max_cells": 4, "max_chars_per_cell": 150},
            "description": "4 ô card, hợp với nhóm phát hiện/nguyên nhân/thành phần.",
        },
        {
            "id": "numbered-columns-4",
            "intent": ["recommend_actions"],
            "supported_blocks": ["title", "columns", "footer_text"],
            "capacity": {"max_columns": 4, "max_bullets_per_column": 3},
            "description": "4 cột đánh số, hợp với quy trình hoặc khuyến nghị hành động.",
        },
        {
            "id": "timeline",
            "intent": ["show_progression"],
            "supported_blocks": ["title", "milestones", "footer_text"],
            "capacity": {"max_milestones": 5},
            "description": "Đường thời gian ngang, hợp với lộ trình và milestone.",
        },
        {
            "id": "stair-progress",
            "intent": ["show_progression"],
            "supported_blocks": ["title", "steps", "footer_text"],
            "capacity": {"max_steps": 5},
            "description": "Bậc thang tiến trình, hợp với maturity/progress.",
        },
    ],
}


DECK_TEMPLATE = r'''
presentation "{{ title }}":
  meta:
    author: "Ontology Task Graph Pipeline"
    language: vi
    format: pptx

  use style: "corporate.psl"

  cover_layout: title-slide
  cover:
    subtitle: "Sinh slide theo task graph: phân rã → hoàn thành → tổng hợp"
    author: "LangGraph + Ontology + Jinja + Renderer"

  section "Tổng hợp báo cáo":
    header:
      subtitle: "Mỗi slide được sinh từ một task độc lập và đã validate"

{% for slide in slides %}
    slide "{{ slide.title | e }}":
      layout: {{ slide.layout }}
      intent: {{ slide.intent }}

      title:
        {{ slide.title | e }}
{% if slide.subtitle %}

      subtitle:
        {{ slide.subtitle | e }}
{% endif %}
{% if slide.blocks.get("bullets") %}

      bullets:
{% if slide.blocks.get("bullets") is mapping %}
{% if slide.blocks.get("bullets").get("icon") %}
        icon: {{ slide.blocks.get("bullets").get("icon") }}
{% endif %}
        items:
{% for item in slide.blocks.get("bullets").get("items", []) %}
          - {{ item | e }}
{% endfor %}
{% else %}
{% for item in slide.blocks.get("bullets") %}
        - {{ item | e }}
{% endfor %}
{% endif %}
{% endif %}
{% if slide.blocks.get("cells") %}

      cells:
{% for cell in slide.blocks.get("cells") %}
        - heading: {{ cell.heading | e }}
          text: {{ cell.text | e }}
{% endfor %}
{% endif %}
{% if slide.blocks.get("columns") %}

      columns:
{% for col in slide.blocks.get("columns") %}
        - heading: {{ col.heading | e }}
          text: {{ col.text | e }}
{% if col.bullets %}
          bullets:
{% for b in col.bullets %}
            - {{ b | e }}
{% endfor %}
{% endif %}
{% endfor %}
{% endif %}
{% if slide.blocks.get("milestones") %}

      milestones:
{% for m in slide.blocks.get("milestones") %}
        - date: {{ m.date | e }}
          heading: {{ m.heading | e }}
          text: {{ m.text | e }}
{% endfor %}
{% endif %}
{% if slide.blocks.get("steps") %}

      steps:
{% for s in slide.blocks.get("steps") %}
        - heading: {{ s.heading | e }}
          text: {{ s.text | e }}
{% endfor %}
{% endif %}
{% if slide.blocks.get("conclusion") %}

      conclusion:
        icon: check
        text: {{ slide.blocks.get("conclusion") | e }}
{% endif %}

      footer_text:
        Task {{ slide.task_id }} | source={{ slide.source_concept_id }} | layout={{ slide.layout }}

{% endfor %}
'''


DEFAULT_PSL = r'''
theme "corporate":
  page:
    size: widescreen
    background: "#FFFFFF"

  colors:
    primary: "#003A8C"
    secondary: "#E6F0FF"
    text: "#1F1F1F"
    muted: "#666666"
    white: "#FFFFFF"

  fonts:
    heading: "Aptos Display"
    body: "Aptos"

  presentation.title-slide:
    background: primary
    title:
      font: heading
      size: 44
      color: white
      bold: true
      align: center
      position: [80, 210]
      width: 1120
      height: 90
    subtitle:
      font: body
      size: 24
      color: secondary
      align: center
      position: [100, 315]
      width: 1080
      height: 60
    author:
      font: body
      size: 18
      color: white
      align: center
      position: [140, 430]
      width: 1000
      height: 40

  section.section-header:
    background: secondary
    title:
      font: heading
      size: 42
      color: primary
      bold: true
      align: center
      position: [90, 250]
      width: 1100
      height: 80
    subtitle:
      font: body
      size: 22
      color: text
      align: center
      position: [120, 345]
      width: 1040
      height: 60

  slide.title-bullets:
    title:
      font: heading
      size: 34
      color: primary
      bold: true
      position: [60, 40]
      width: 1120
      height: 70
    subtitle:
      font: body
      size: 19
      color: muted
      position: [65, 105]
      width: 1080
      height: 38
    bullets:
      font: body
      size: 22
      color: text
      position: [85, 165]
      width: 1080
      height: 340
      line_gap: 10
      overflow: shrink
      min_size: 15
    conclusion:
      font: body
      size: 18
      color: text
      bold: true
      align: center
      fill: "#E6F0FF"
      border: primary
      position: [90, 555]
      width: 1100
      height: 80
    footer_text:
      font: body
      size: 12
      color: muted
      position: [70, 685]
      width: 900
      height: 24

  slide.grid-4:
    title:
      font: heading
      size: 32
      color: primary
      bold: true
      position: [60, 40]
      width: 1120
      height: 70
    grid:
      position: [70, 150]
      width: 1140
      height: 430
      columns: 2
      gap: 22
    card:
      fill: "#F8FAFC"
      border: "#CBD5E1"
    footer_text:
      font: body
      size: 12
      color: muted
      position: [70, 685]
      width: 900
      height: 24

  slide.numbered-columns-4:
    title:
      font: heading
      size: 32
      color: primary
      bold: true
      position: [60, 40]
      width: 1120
      height: 70
    columns:
      position: [70, 165]
      width: 1140
      height: 410
      gap: 20
    card:
      fill: "#FFFFFF"
      border: "#CBD5E1"
    footer_text:
      font: body
      size: 12
      color: muted
      position: [70, 685]
      width: 900
      height: 24

  slide.timeline:
    title:
      font: heading
      size: 32
      color: primary
      bold: true
      position: [60, 40]
      width: 1120
      height: 70
    timeline:
      position: [90, 180]
      width: 1100
      height: 420
      line_color: primary
      card_fill: "#F8FAFC"
      card_border: "#CBD5E1"
    footer_text:
      font: body
      size: 12
      color: muted
      position: [70, 685]
      width: 900
      height: 24

  slide.stair-progress:
    title:
      font: heading
      size: 32
      color: primary
      bold: true
      position: [60, 40]
      width: 1120
      height: 70
    stair:
      position: [90, 170]
      width: 1080
      height: 430
      fill: "#F8FAFC"
      border: "#CBD5E1"
    footer_text:
      font: body
      size: 12
      color: muted
      position: [70, 685]
      width: 900
      height: 24
'''


def ensure_list(value: Any) -> List[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def safe_text(value: Any, fallback: str = "") -> str:
    if value is None:
        return fallback
    text = str(value).strip()
    return text if text else fallback


def truncate(text: str, limit: int) -> str:
    text = safe_text(text)
    return text if len(text) <= limit else text[: limit - 1].rstrip() + "…"


def load_json_or_default(path: Optional[str], default: Dict[str, Any], out_dir: Path, filename: str) -> Dict[str, Any]:
    if path:
        p = Path(path)
        if p.exists():
            return json.loads(p.read_text(encoding="utf-8"))
    out_path = out_dir / filename
    out_path.write_text(json.dumps(default, ensure_ascii=False, indent=2), encoding="utf-8")
    return deepcopy(default)


def extract_docx_text(docx_path: Path) -> str:
    """Extract readable text from a .docx file.

    The extractor keeps paragraph order and also serializes tables in a
    Markdown-like pipe format so the downstream LLM can still understand
    tabular report content.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Missing dependency: python-docx. Install with: pip install python-docx") from exc

    doc = Document(str(docx_path))
    parts: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style_name = getattr(para.style, "name", "") if para.style else ""
            # Preserve heading-like cues for ontology-based summarization.
            if style_name.lower().startswith("heading"):
                parts.append(f"# {text}")
            else:
                parts.append(text)

    for table_index, table in enumerate(doc.tables, start=1):
        rows: List[List[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            parts.append(f"\n[Table {table_index}]")
            parts.extend(_format_rows_as_markdown_table(rows))

    return "\n\n".join(parts).strip()


def _format_rows_as_markdown_table(rows: List[List[str]]) -> List[str]:
    if not rows:
        return []
    width = max(len(r) for r in rows)
    normalized = [r + [""] * (width - len(r)) for r in rows]
    lines = []
    header = normalized[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def extract_pdf_text(pdf_path: Path) -> str:
    try:
        import pypdf  # type: ignore
    except ImportError as exc:
        raise RuntimeError("For local PDF text extraction, install pypdf or use --mock / Gemini File API.") from exc
    reader = pypdf.PdfReader(str(pdf_path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def extract_plain_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def normalize_raw_input_text(text: str) -> str:
    """Normalize direct plain-text input while preserving report structure.

    This deliberately does not over-clean the text: administrative reports often
    use headings, numbered sections, colon-separated labels, and bullet-like
    lines that are useful signals for the ontology-guided summarizer.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def extract_input_text(input_path: Path, mock: bool, direct_text: str = "", use_stdin: bool = False) -> str:
    """Extract text from PDF, DOCX, TXT/MD, direct --text, or stdin.

    Precedence:
      1. --mock sample text
      2. --text direct plain text
      3. --stdin content
      4. input_path file extraction
    """
    if mock:
        return (
            "Báo cáo nêu vấn đề quá tải thông tin trong doanh nghiệp. "
            "Dữ liệu phân tán ở nhiều hệ thống như CRM, ERP, báo cáo thủ công và mạng xã hội. "
            "Các phát hiện chính gồm thiếu metadata thống nhất, quy trình tổng hợp chậm, "
            "và khó truy xuất bằng ngữ cảnh. Rủi ro gồm sai lệch số liệu, trễ báo cáo, "
            "và phụ thuộc chuyên gia thủ công. Khuyến nghị là xây dựng pipeline tổng hợp tri thức, "
            "chuẩn hóa dữ liệu, dùng RAG/ontology, tự động sinh dashboard và slide. "
            "Lộ trình gồm Q1 khảo sát, Q2 chuẩn hóa, Q3 triển khai thử nghiệm, Q4 mở rộng."
        )

    if direct_text:
        return normalize_raw_input_text(direct_text)

    if use_stdin:
        stdin_text = sys.stdin.read()
        if not stdin_text.strip():
            raise ValueError("--stdin was provided but stdin is empty.")
        return normalize_raw_input_text(stdin_text)

    if not input_path.exists():
        raise FileNotFoundError(
            f"Input file not found: {input_path}. "
            "Use --text '...' for direct plain text or --stdin to pipe content."
        )

    suffix = input_path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(input_path)
    if suffix == ".docx":
        return extract_docx_text(input_path)
    if suffix in {".txt", ".md", ".markdown"}:
        return extract_plain_text(input_path)

    raise ValueError(
        f"Unsupported input format: {suffix}. Supported formats: .pdf, .docx, .txt, .md"
    )


def call_gemini_json(prompt: str, model: str, mock_json: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    if mock_json is not None:
        return mock_json
    try:
        from google import genai  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Missing google-genai. Install with: pip install -U google-genai") from exc
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY is not set.")
    client = genai.Client(api_key=api_key)
    resp = client.models.generate_content(model=model, contents=prompt)
    raw = getattr(resp, "text", "") or ""
    return parse_json_from_text(raw)


def parse_json_from_text(raw: str) -> Dict[str, Any]:
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not m:
            raise
        return json.loads(m.group(0))


def validate_input(state: PipelineState) -> PipelineState:
    out_dir = Path(state["out_dir"])
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "task_outputs").mkdir(exist_ok=True)
    state.setdefault("validation_reports", [])
    state.setdefault("repair_log", [])
    return state


def load_planning_assets(state: PipelineState) -> PipelineState:
    out_dir = Path(state["out_dir"])
    state["report_ontology"] = load_json_or_default(
        state.get("report_ontology_path"), DEFAULT_REPORT_ONTOLOGY, out_dir, "report_domain_ontology.json"
    )
    state["slide_ontology"] = load_json_or_default(
        state.get("slide_ontology_path"), DEFAULT_SLIDE_ONTOLOGY, out_dir, "slide_creation_ontology.json"
    )
    state["layout_registry"] = load_json_or_default(
        state.get("layout_registry_path"), DEFAULT_LAYOUT_REGISTRY, out_dir, "layout_registry.json"
    )
    return state


def summarize_report_by_ontology(state: PipelineState) -> PipelineState:
    source_text = extract_input_text(
        Path(state["pdf_path"]),
        state.get("mock", False),
        state.get("input_text", ""),
        state.get("use_stdin", False),
    )
    state["source_text"] = source_text
    concepts = state["report_ontology"]["concepts"]

    mock_summary = {
        "title": "Báo cáo tổng hợp tri thức doanh nghiệp",
        "sections": [
            {
                "concept_id": c["id"],
                "heading": c["preferred_heading"],
                "intent": c["intent"],
                "bullets": make_mock_bullets(c["id"]),
                "conclusion": make_mock_conclusion(c["id"]),
            }
            for c in concepts
        ],
    }

    prompt = f"""
Bạn là chuyên gia phân tích báo cáo.
Hãy tóm tắt văn bản theo ontology nghiệp vụ sau, trả về JSON thuần.

Ontology concepts:
{json.dumps(concepts, ensure_ascii=False, indent=2)}

Schema:
{{
  "title": "...",
  "sections": [
    {{"concept_id": "...", "heading": "...", "intent": "...", "bullets": ["..."], "conclusion": "..."}}
  ]
}}

Văn bản:
{source_text[:30000]}
"""
    summary = call_gemini_json(prompt, state.get("model", "gemini-2.0-flash"), mock_summary if state.get("mock") else None)
    state["report_summary"] = normalize_summary(summary, state["report_ontology"])
    return state


def make_mock_bullets(concept_id: str) -> List[str]:
    mapping = {
        "executive_summary": [
            "Doanh nghiệp đang gặp tình trạng quá tải thông tin và thiếu cơ chế tổng hợp nhanh.",
            "Giá trị chính của hệ thống là chuyển dữ liệu phân tán thành tri thức có thể trình bày.",
            "Cần kết hợp chuẩn hóa dữ liệu, truy xuất ngữ cảnh và sinh báo cáo tự động.",
        ],
        "context": [
            "Dữ liệu nằm rải rác ở nhiều hệ thống nghiệp vụ.",
            "Báo cáo thủ công làm chậm quá trình ra quyết định.",
            "Lãnh đạo cần insight nhanh hơn, nhất quán hơn và truy vết được nguồn.",
        ],
        "key_findings": [
            "Thiếu metadata thống nhất khiến truy xuất khó ổn định.",
            "Chunking và tổng hợp thủ công dễ làm mất ngữ cảnh.",
            "Pipeline hiện tại chưa tự động hóa đủ các bước từ dữ liệu đến slide.",
            "Ontology có thể giúp chuẩn hóa tiêu đề và cấu trúc báo cáo.",
        ],
        "risks": [
            "Rủi ro sai lệch số liệu nếu nguồn dữ liệu không được kiểm soát.",
            "Rủi ro hallucination nếu LLM thiếu grounding và validation.",
            "Rủi ro trễ tiến độ nếu không chia pipeline thành các task nhỏ.",
        ],
        "recommendations": [
            "Chuẩn hóa ontology nghiệp vụ báo cáo trước khi tóm tắt.",
            "Dùng layout registry để chọn bố cục theo intent thay vì hardcode.",
            "Thêm validate/repair ở từng task trước khi tổng hợp deck.",
            "Tách sinh slide thành nhiều nhiệm vụ độc lập trong LangGraph.",
        ],
        "roadmap": ["Q1: khảo sát nguồn dữ liệu", "Q2: chuẩn hóa ontology", "Q3: thử nghiệm pipeline", "Q4: mở rộng triển khai"],
    }
    return mapping.get(concept_id, ["Nội dung chính cần được tổng hợp theo ontology."])


def make_mock_conclusion(concept_id: str) -> str:
    mapping = {
        "executive_summary": "Trọng tâm là biến báo cáo dài thành slide có cấu trúc, kiểm soát được và dễ trình bày.",
        "context": "Bài toán không chỉ là thiếu dữ liệu, mà là thiếu năng lực tổ chức và khai thác dữ liệu.",
        "key_findings": "Các phát hiện cho thấy cần một pipeline có ontology, registry và validation rõ ràng.",
        "risks": "Rủi ro phải được kiểm soát bằng validation, grounding và cơ chế repair tự động.",
        "recommendations": "Nên triển khai theo từng task nhỏ để dễ kiểm soát chất lượng.",
        "roadmap": "Lộ trình nên đi từ chuẩn hóa nền tảng đến thử nghiệm và mở rộng.",
    }
    return mapping.get(concept_id, "Cần chuyển nội dung thành hành động rõ ràng.")


def normalize_summary(summary: Dict[str, Any], ontology: Dict[str, Any]) -> Dict[str, Any]:
    concept_map = {c["id"]: c for c in ontology.get("concepts", [])}
    sections_by_id = {s.get("concept_id"): s for s in ensure_list(summary.get("sections")) if isinstance(s, dict)}
    normalized_sections: List[Dict[str, Any]] = []
    for concept in ontology.get("concepts", []):
        cid = concept["id"]
        src = sections_by_id.get(cid, {})
        bullets = [truncate(safe_text(b), 180) for b in ensure_list(src.get("bullets"))]
        if not bullets:
            bullets = ["Chưa trích xuất được nội dung; cần kiểm tra nguồn báo cáo."]
        normalized_sections.append(
            {
                "concept_id": cid,
                "heading": safe_text(src.get("heading"), concept.get("preferred_heading", cid)),
                "intent": safe_text(src.get("intent"), concept.get("intent", "summarize_core_message")),
                "bullets": bullets[: int(concept.get("max_bullets", 5))],
                "conclusion": truncate(safe_text(src.get("conclusion"), "Cần tiếp tục rà soát và triển khai theo ưu tiên."), 220),
            }
        )
    return {"title": safe_text(summary.get("title"), "Báo cáo tổng hợp"), "sections": normalized_sections}


def choose_layout(intent: str, section: Dict[str, Any], slide_ontology: Dict[str, Any], registry: Dict[str, Any]) -> str:
    available = {layout["id"]: layout for layout in registry.get("layouts", [])}
    for rule in slide_ontology.get("layout_selection_rules", []):
        if rule.get("intent") == intent:
            for candidate in rule.get("prefer", []):
                if candidate in available:
                    return candidate
    for layout in registry.get("layouts", []):
        if intent in layout.get("intent", []):
            return layout["id"]
    return "title-bullets"


def decompose_into_slide_tasks(state: PipelineState) -> PipelineState:
    summary = state["report_summary"]
    tasks: List[Dict[str, Any]] = []
    for idx, section in enumerate(summary.get("sections", []), start=1):
        layout_id = choose_layout(section.get("intent", ""), section, state["slide_ontology"], state["layout_registry"])
        task = {
            "task_id": f"T{idx:02d}",
            "task_type": "create_slide_fragment",
            "source_concept_id": section["concept_id"],
            "slide_intent": section.get("intent", "summarize_core_message"),
            "layout_id": layout_id,
            "title": section.get("heading"),
            "input_summary_section": section,
            "status": "pending",
        }
        tasks.append(task)

    state["task_plan"] = {
        "deck_title": summary.get("title", "Báo cáo tổng hợp"),
        "policy": state["slide_ontology"].get("task_policy", {}),
        "tasks": tasks,
    }
    state["pending_tasks"] = tasks
    state["completed_tasks"] = []
    return state


def has_pending_tasks(state: PipelineState) -> str:
    return "process_next_task" if state.get("pending_tasks") else "aggregate_deck"


def pop_next_task(state: PipelineState) -> PipelineState:
    pending = list(state.get("pending_tasks", []))
    if not pending:
        return state
    state["current_task"] = pending.pop(0)
    state["pending_tasks"] = pending
    return state


def process_current_task(state: PipelineState) -> PipelineState:
    task = state["current_task"]
    section = task["input_summary_section"]
    layout = task["layout_id"]

    # Each task produces a structured slide spec, not PML directly.
    if layout == "grid-4":
        cells = []
        for b in section["bullets"][:4]:
            parts = b.split(":", 1)
            heading = parts[0] if len(parts) > 1 and len(parts[0]) < 30 else "Ý chính"
            text = parts[1].strip() if len(parts) > 1 else b
            cells.append({"heading": truncate(heading, 45), "text": truncate(text, 150)})
        blocks: Dict[str, Any] = {"cells": cells, "conclusion": section.get("conclusion")}
    elif layout == "numbered-columns-4":
        columns = []
        for b in section["bullets"][:4]:
            columns.append({"heading": truncate(b, 45), "text": truncate(section.get("conclusion", ""), 120), "bullets": []})
        blocks = {"columns": columns}
    elif layout == "timeline":
        milestones = []
        for i, b in enumerate(section["bullets"][:5], start=1):
            label = f"M{i}"
            if re.match(r"^(Q\d|20\d{2}|Giai đoạn|Phase)", b):
                label = b.split(":", 1)[0][:12]
            milestones.append({"date": label, "heading": truncate(b.split(":", 1)[0], 45), "text": truncate(b, 150)})
        blocks = {"milestones": milestones}
    elif layout == "stair-progress":
        blocks = {"steps": [{"heading": truncate(b, 50), "text": truncate(b, 150)} for b in section["bullets"][:5]]}
    else:
        blocks = {
            "bullets": {"icon": "check", "items": section["bullets"]},
            "conclusion": section.get("conclusion"),
        }

    state["current_task_output"] = {
        "task_id": task["task_id"],
        "source_concept_id": task["source_concept_id"],
        "intent": task["slide_intent"],
        "title": task["title"],
        "subtitle": f"Nguồn ontology: {task['source_concept_id']}",
        "layout": layout,
        "blocks": blocks,
        "status": "completed",
    }
    return state


def validate_current_task_output(state: PipelineState) -> PipelineState:
    output = state["current_task_output"]
    registry = {x["id"]: x for x in state["layout_registry"].get("layouts", [])}
    issues = []
    if not output.get("title"):
        issues.append({"severity": "error", "field": "title", "message": "Missing slide title"})
    if output.get("layout") not in registry:
        issues.append({"severity": "error", "field": "layout", "message": f"Unsupported layout {output.get('layout')}"})
    else:
        supported = set(registry[output["layout"]].get("supported_blocks", []))
        for block_name in output.get("blocks", {}).keys():
            if block_name not in supported and block_name != "conclusion":
                issues.append({"severity": "warning", "field": block_name, "message": "Block may not be supported by layout"})

    output["validation_issues"] = issues
    state.setdefault("validation_reports", []).append(
        {"stage": "validate_task_output", "task_id": output.get("task_id"), "issues": issues}
    )
    return state


def repair_current_task_output(state: PipelineState) -> PipelineState:
    output = state["current_task_output"]
    issues = output.get("validation_issues", [])
    if not issues:
        return state

    registry_ids = {x["id"] for x in state["layout_registry"].get("layouts", [])}
    if not output.get("title"):
        output["title"] = "Slide chưa đặt tiêu đề"
        state.setdefault("repair_log", []).append(f"{output.get('task_id')}: added fallback title")
    if output.get("layout") not in registry_ids:
        output["layout"] = "title-bullets"
        output["blocks"] = {"bullets": ["Nội dung đã được chuyển về layout an toàn."]}
        state.setdefault("repair_log", []).append(f"{output.get('task_id')}: changed unsupported layout to title-bullets")
    output["validation_issues"] = []
    return state


def store_completed_task(state: PipelineState) -> PipelineState:
    output = state["current_task_output"]
    completed = list(state.get("completed_tasks", []))
    completed.append(output)
    state["completed_tasks"] = completed

    out_dir = Path(state["out_dir"])
    task_path = out_dir / "task_outputs" / f"{output['task_id']}.json"
    task_path.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
    return state


def aggregate_deck(state: PipelineState) -> PipelineState:
    env = Environment(undefined=StrictUndefined, trim_blocks=False, lstrip_blocks=False)
    template = env.from_string(DECK_TEMPLATE)
    slides = sorted(state.get("completed_tasks", []), key=lambda x: x.get("task_id", ""))
    state["pml_text"] = template.render(title=state["task_plan"]["deck_title"], slides=slides)
    state["psl_text"] = DEFAULT_PSL
    return state


def validate_aggregated_pml(state: PipelineState) -> PipelineState:
    pml = state.get("pml_text", "")
    issues = []
    if not pml.strip().startswith("presentation"):
        issues.append({"severity": "error", "message": "PML does not start with presentation"})
    if "section " not in pml:
        issues.append({"severity": "error", "message": "PML has no section"})
    if "slide " not in pml:
        issues.append({"severity": "error", "message": "PML has no slide"})
    state.setdefault("validation_reports", []).append({"stage": "validate_aggregated_pml", "issues": issues})
    if issues:
        state.setdefault("repair_log", []).append("Aggregated PML had issues; applying fallback repair.")
        state["pml_text"] = fallback_pml(state)
    return state


def fallback_pml(state: PipelineState) -> str:
    return '''presentation "Fallback Deck":
  use style: "corporate.psl"
  section "Fallback":
    slide "Không tạo được deck đầy đủ":
      layout: title-bullets
      title:
        Không tạo được deck đầy đủ
      bullets:
        - Pipeline đã fallback do lỗi validation.
'''


def write_outputs(state: PipelineState) -> PipelineState:
    out_dir = Path(state["out_dir"])
    pml_path = out_dir / "generated.pml"
    psl_path = out_dir / "corporate.psl"
    task_plan_path = out_dir / "task_plan.json"
    validation_path = out_dir / "validation_report.json"
    repair_path = out_dir / "repair_log.txt"
    source_text_path = out_dir / "source_text.txt"

    source_text_path.write_text(state.get("source_text", ""), encoding="utf-8")
    pml_path.write_text(state["pml_text"], encoding="utf-8")
    psl_path.write_text(state["psl_text"], encoding="utf-8")
    task_plan_path.write_text(json.dumps(state["task_plan"], ensure_ascii=False, indent=2), encoding="utf-8")
    validation_path.write_text(json.dumps(state.get("validation_reports", []), ensure_ascii=False, indent=2), encoding="utf-8")
    repair_path.write_text("\n".join(state.get("repair_log", [])), encoding="utf-8")

    state["pml_path"] = str(pml_path)
    state["psl_path"] = str(psl_path)
    state["task_plan_path"] = str(task_plan_path)
    state["source_text_path"] = str(source_text_path)
    return state


def render_with_existing_renderer(state: PipelineState) -> PipelineState:
    renderer_path = Path(state["renderer_path"])
    if not renderer_path.exists():
        # Try relative to current working directory or /mnt/data.
        alt = Path.cwd() / state["renderer_path"]
        if alt.exists():
            renderer_path = alt
        else:
            alt2 = Path("/mnt/data") / state["renderer_path"]
            if alt2.exists():
                renderer_path = alt2
    if not renderer_path.exists():
        state.setdefault("repair_log", []).append(f"Renderer not found: {state['renderer_path']}; skipped render.")
        return state

    spec = importlib.util.spec_from_file_location("dsl_renderer", str(renderer_path))
    if spec is None or spec.loader is None:
        state.setdefault("repair_log", []).append("Could not import renderer; skipped render.")
        return state
    module = importlib.util.module_from_spec(spec)
    sys.modules["dsl_renderer"] = module
    spec.loader.exec_module(module)  # type: ignore

    required = ["parse_pml", "parse_psl", "build_render_ir", "render_html", "render_pptx"]
    missing = [name for name in required if not hasattr(module, name)]
    if missing:
        state.setdefault("repair_log", []).append(f"Renderer missing functions {missing}; skipped render.")
        return state

    doc = module.parse_pml(state["pml_text"])
    theme = module.parse_psl(state["psl_text"])
    render_ir = module.build_render_ir(doc, theme)

    out_dir = Path(state["out_dir"])
    html_path = out_dir / "output.html"
    pptx_path = out_dir / "output.pptx"
    md_path = out_dir / "output.md"
    module.render_html(render_ir, str(html_path))
    module.render_pptx(render_ir, str(pptx_path))
    if hasattr(module, "render_markdown"):
        module.render_markdown(render_ir, str(md_path))
        state["md_path"] = str(md_path)
    state["html_path"] = str(html_path)
    state["pptx_path"] = str(pptx_path)
    return state


def build_graph():
    END, START, StateGraph = require_langgraph()
    graph = StateGraph(PipelineState)
    graph.add_node("validate_input", validate_input)
    graph.add_node("load_planning_assets", load_planning_assets)
    graph.add_node("summarize_report_by_ontology", summarize_report_by_ontology)
    graph.add_node("decompose_into_slide_tasks", decompose_into_slide_tasks)
    graph.add_node("pop_next_task", pop_next_task)
    graph.add_node("process_current_task", process_current_task)
    graph.add_node("validate_current_task_output", validate_current_task_output)
    graph.add_node("repair_current_task_output", repair_current_task_output)
    graph.add_node("store_completed_task", store_completed_task)
    graph.add_node("aggregate_deck", aggregate_deck)
    graph.add_node("validate_aggregated_pml", validate_aggregated_pml)
    graph.add_node("write_outputs", write_outputs)
    graph.add_node("render_with_existing_renderer", render_with_existing_renderer)

    graph.add_edge(START, "validate_input")
    graph.add_edge("validate_input", "load_planning_assets")
    graph.add_edge("load_planning_assets", "summarize_report_by_ontology")
    graph.add_edge("summarize_report_by_ontology", "decompose_into_slide_tasks")
    graph.add_conditional_edges(
        "decompose_into_slide_tasks",
        has_pending_tasks,
        {"process_next_task": "pop_next_task", "aggregate_deck": "aggregate_deck"},
    )
    graph.add_edge("pop_next_task", "process_current_task")
    graph.add_edge("process_current_task", "validate_current_task_output")
    graph.add_edge("validate_current_task_output", "repair_current_task_output")
    graph.add_edge("repair_current_task_output", "store_completed_task")
    graph.add_conditional_edges(
        "store_completed_task",
        has_pending_tasks,
        {"process_next_task": "pop_next_task", "aggregate_deck": "aggregate_deck"},
    )
    graph.add_edge("aggregate_deck", "validate_aggregated_pml")
    graph.add_edge("validate_aggregated_pml", "write_outputs")
    graph.add_edge("write_outputs", "render_with_existing_renderer")
    graph.add_edge("render_with_existing_renderer", END)
    return graph.compile()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Task-graph ontology slide generation pipeline for PDF/DOCX/TXT/direct plain-text input")
    parser.add_argument("input_path", nargs="?", default="", help="Input .pdf, .docx, .txt, or .md path. Optional when using --text, --stdin, or --mock.")
    parser.add_argument("--text", default="", help="Direct plain-text input. Use this when content is already available as text, not as a file.")
    parser.add_argument("--stdin", action="store_true", help="Read plain-text input from stdin. Example: cat report.txt | python script.py --stdin")
    parser.add_argument("--renderer", default="demo_dsl_conclusion_box.py", help="Path to renderer .py file")
    parser.add_argument("--out-dir", default="out_task_graph_demo", help="Output directory")
    parser.add_argument("--model", default="gemini-2.0-flash", help="Gemini model")
    parser.add_argument("--mock", action="store_true", help="Run without Gemini/PDF extraction")
    parser.add_argument("--report-ontology", default="", help="Optional report ontology JSON path")
    parser.add_argument("--slide-ontology", default="", help="Optional slide ontology JSON path")
    parser.add_argument("--layout-registry", default="", help="Optional layout registry JSON path")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    app = build_graph()
    initial: PipelineState = {
        "pdf_path": args.input_path or "direct_text_input.txt",
        "input_text": args.text,
        "use_stdin": args.stdin,
        "renderer_path": args.renderer,
        "out_dir": args.out_dir,
        "model": args.model,
        "mock": args.mock,
        "report_ontology_path": args.report_ontology,
        "slide_ontology_path": args.slide_ontology,
        "layout_registry_path": args.layout_registry,
    }
    final = app.invoke(initial)
    print("Generated:")
    for key in ["source_text_path", "pml_path", "psl_path", "task_plan_path", "html_path", "pptx_path", "md_path"]:
        if final.get(key):
            print(f"- {key}: {final[key]}")
    print(f"- task_outputs: {Path(args.out_dir) / 'task_outputs'}")
    print(f"- validation_report: {Path(args.out_dir) / 'validation_report.json'}")
    print(f"- repair_log: {Path(args.out_dir) / 'repair_log.txt'}")


if __name__ == "__main__":
    main()
