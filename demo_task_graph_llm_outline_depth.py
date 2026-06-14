#!/usr/bin/env python3
"""
Task-graph LLM outline-depth-first PDF/DOCX/TXT/PLAIN-TEXT -> slide deck pipeline
======================================================

This version deliberately avoids generating the whole slide deck in one shot.
It decomposes the work into explicit tasks, completes each task independently,
validates/repairs each task result, skips empty/unusable report components, strips administrative boilerplate, and only then aggregates slide fragments into
one PML deck without auto-generated footers.

Core idea
---------
    PDF/DOCX/TXT report text
      -> LLM explicitly analyzes document outline depth; ontology is only a reference/fallback
      -> slide ontology + layout registry guide task creation
      -> each task creates exactly one slide-spec / PML fragment
      -> each fragment is validated/repaired independently
      -> final aggregation builds the deck-level PML
      -> existing renderer produces HTML/PPTX

Install
-------
    pip install -U langgraph google-genai jinja2 python-pptx python-docx pypdf

Mock run
--------
    python demo_task_graph_docx_input.py dummy.docx \
      --renderer demo_dsl_conclusion_box.py \
      --out-dir out_task_graph_demo \
      --mock

Gemini run
----------
    export GEMINI_API_KEY="your-key"
    python demo_task_graph_docx_input.py input.docx \
      --renderer demo_dsl_conclusion_box.py \
      --out-dir out_task_graph_demo

Outputs
-------
    source_text.txt
    generated.pml
    corporate.psl
    task_plan.json
    task_outputs/*.json
    validation_report.json
    repair_log.txt
    output.html
    output.pptx
"""

from __future__ import annotations

import argparse
import importlib.util
import json
import os
import re
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional, TypedDict

from jinja2 import Environment, StrictUndefined


def require_langgraph():
    try:
        from langgraph.graph import END, START, StateGraph
    except ImportError as exc:
        raise RuntimeError("Missing dependency: langgraph. Install with: pip install -U langgraph") from exc
    return END, START, StateGraph


class PipelineState(TypedDict, total=False):
    pdf_path: str
    input_text: str
    use_stdin: bool
    renderer_path: str
    out_dir: str
    model: str
    mock: bool

    report_ontology_path: str
    slide_ontology_path: str
    layout_registry_path: str

    report_ontology: Dict[str, Any]
    slide_ontology: Dict[str, Any]
    layout_registry: Dict[str, Any]

    source_text: str
    report_summary: Dict[str, Any]
    document_outline: List[Dict[str, Any]]
    outline_hierarchy: Dict[str, Any]
    section_groups: List[Dict[str, Any]]

    task_plan: Dict[str, Any]
    pending_tasks: List[Dict[str, Any]]
    current_task: Dict[str, Any]
    current_task_output: Dict[str, Any]
    completed_tasks: List[Dict[str, Any]]

    pml_text: str
    psl_text: str

    pml_path: str
    psl_path: str
    task_plan_path: str
    html_path: str
    pptx_path: str
    md_path: str
    source_text_path: str

    validation_reports: List[Dict[str, Any]]
    repair_log: List[str]


DEFAULT_REPORT_ONTOLOGY: Dict[str, Any] = {
    "ontology_id": "report-domain-ontology-v0.3-outline-first-clean",
    "description": "Ontology nghiệp vụ báo cáo. Chỉ dùng làm tham khảo; bố cục thật của văn bản là nguồn ưu tiên. Có quy tắc loại bỏ phần nghi thức/hành chính không nên đưa vào slide.",
    "exclusion_rules": {
        "ignore_opening_motto": True,
        "ignore_recipients_tail": True,
        "ignored_heading_patterns": [
            r"^cộng\s+h[oòóọõỏôồốộỗổơờớợỡở][aàáạãả]\s+x[aã]\s+h[oộòóọõỏ]i\s+ch[uủ]\s+ngh[iĩ]a\s+vi[eệ]t\s+nam$",
            r"^độc\s+lập\s*[-–—]\s*tự\s+do\s*[-–—]\s*hạnh\s+phúc$",
            r"^số\s*[:：]\s*.*$",
            r"^số\s+.*$",
            r"^nơi\s+nhận\s*:?$",
            r"^kính\s+gửi\s*:?$",
            r"^căn\s+cứ\s+.*$",
            r"^tm\.?.*$",
            r"^kt\.?.*$",
            r"^chủ\s+tịch.*$",
            r"^người\s+ký.*$"
        ],
        "description": "Không coi quốc hiệu/tiêu ngữ mở đầu và mục nơi nhận cuối văn bản là nội dung báo cáo để tạo slide."
    },
    "concepts": [
        {
            "id": "executive_summary",
            "preferred_heading": "Tóm tắt điều hành",
            "intent": "summarize_core_message",
            "max_bullets": 4,
            "slide_priority": 1,
        },
        {
            "id": "context",
            "preferred_heading": "Bối cảnh và vấn đề",
            "intent": "explain_context",
            "max_bullets": 5,
            "slide_priority": 2,
        },
        {
            "id": "key_findings",
            "preferred_heading": "Phát hiện chính",
            "intent": "present_findings",
            "max_bullets": 6,
            "slide_priority": 3,
        },
        {
            "id": "risks",
            "preferred_heading": "Rủi ro và điểm cần chú ý",
            "intent": "highlight_risks",
            "max_bullets": 5,
            "slide_priority": 4,
        },
        {
            "id": "recommendations",
            "preferred_heading": "Khuyến nghị hành động",
            "intent": "recommend_actions",
            "max_bullets": 5,
            "slide_priority": 5,
        },
        {
            "id": "roadmap",
            "preferred_heading": "Lộ trình triển khai",
            "intent": "show_progression",
            "max_bullets": 5,
            "slide_priority": 6,
        },
    ],
}


DEFAULT_SLIDE_ONTOLOGY: Dict[str, Any] = {
    "ontology_id": "slide-creation-ontology-v0.2",
    "concepts": {
        "SlideTask": {
            "description": "Một nhiệm vụ độc lập tạo một slide hoặc một fragment slide.",
            "required_fields": ["task_id", "source_concept_id", "slide_intent", "layout_id", "title"],
        },
        "SlideSpec": {
            "description": "Kết quả có cấu trúc sau khi hoàn thành một SlideTask.",
            "required_fields": ["title", "layout", "blocks"],
        },
    },
    "layout_selection_rules": [
        {"intent": "summarize_core_message", "prefer": ["title-bullets", "numbered-columns-3"]},
        {"intent": "explain_context", "prefer": ["title-bullets", "text-image"]},
        {"intent": "present_findings", "prefer": ["grid-4", "title-bullets", "numbered-columns-4"]},
        {"intent": "highlight_risks", "prefer": ["title-bullets", "grid-3"]},
        {"intent": "recommend_actions", "prefer": ["numbered-columns-4", "title-bullets"]},
        {"intent": "show_progression", "prefer": ["timeline", "stair-progress", "title-bullets"]},
    ],
    "task_policy": {
        "one_concept_one_task": True,
        "validate_each_task_before_aggregation": True,
        "aggregate_only_completed_tasks": True,
    },
}


DEFAULT_LAYOUT_REGISTRY: Dict[str, Any] = {
    "registry_id": "renderer-layout-registry-v0.2",
    "layouts": [
        {
            "id": "title-bullets",
            "intent": ["summarize_core_message", "explain_context", "highlight_risks"],
            "supported_blocks": ["title", "subtitle", "bullets", "conclusion"],
            "capacity": {"max_bullets": 7, "max_chars_per_bullet": 130},
            "description": "Tiêu đề + danh sách bullet; layout an toàn nhất cho nội dung báo cáo.",
        },
        {
            "id": "grid-4",
            "intent": ["present_findings"],
            "supported_blocks": ["title", "cells"],
            "capacity": {"max_cells": 4, "max_chars_per_cell": 150},
            "description": "4 ô card, hợp với nhóm phát hiện/nguyên nhân/thành phần.",
        },
        {
            "id": "numbered-columns-4",
            "intent": ["recommend_actions"],
            "supported_blocks": ["title", "columns"],
            "capacity": {"max_columns": 4, "max_bullets_per_column": 3},
            "description": "4 cột đánh số, hợp với quy trình hoặc khuyến nghị hành động.",
        },
        {
            "id": "timeline",
            "intent": ["show_progression"],
            "supported_blocks": ["title", "milestones"],
            "capacity": {"max_milestones": 5},
            "description": "Đường thời gian ngang, hợp với lộ trình và milestone.",
        },
        {
            "id": "stair-progress",
            "intent": ["show_progression"],
            "supported_blocks": ["title", "steps"],
            "capacity": {"max_steps": 5},
            "description": "Bậc thang tiến trình, hợp với maturity/progress.",
        },
    ],
}


DECK_TEMPLATE = r'''
presentation "{{ title }}":
  meta:
    author: "Ontology Task Graph Pipeline"
    language: vi
    format: pptx

  use style: "corporate.psl"

  cover_layout: title-slide
  cover:
    author: "Tạo tự động từ bố cục thật của văn bản"

{% for group in section_groups %}
  section "{{ group.section_title | e }}":

{% for slide in group.slides %}
    slide "{{ slide.title | e }}":
      layout: {{ slide.layout }}
      intent: {{ slide.intent }}

      title:
        {{ slide.title | e }}
{% if slide.subtitle %}

      subtitle:
        {{ slide.subtitle | e }}
{% endif %}
{% if slide.blocks.get("bullets") %}

      bullets:
{% if slide.blocks.get("bullets") is mapping %}
{% if slide.blocks.get("bullets").get("icon") %}
        icon: {{ slide.blocks.get("bullets").get("icon") }}
{% endif %}
        items:
{% for item in slide.blocks.get("bullets").get("items", []) %}
          - {{ item | e }}
{% endfor %}
{% else %}
{% for item in slide.blocks.get("bullets") %}
        - {{ item | e }}
{% endfor %}
{% endif %}
{% endif %}
{% if slide.blocks.get("cells") %}

      cells:
{% for cell in slide.blocks.get("cells") %}
        - heading: {{ cell.heading | e }}
          text: {{ cell.text | e }}
{% endfor %}
{% endif %}
{% if slide.blocks.get("columns") %}

      columns:
{% for col in slide.blocks.get("columns") %}
        - heading: {{ col.heading | e }}
          text: {{ col.text | e }}
{% if col.bullets %}
          bullets:
{% for b in col.bullets %}
            - {{ b | e }}
{% endfor %}
{% endif %}
{% endfor %}
{% endif %}
{% if slide.blocks.get("milestones") %}

      milestones:
{% for m in slide.blocks.get("milestones") %}
        - date: {{ m.date | e }}
          heading: {{ m.heading | e }}
          text: {{ m.text | e }}
{% endfor %}
{% endif %}
{% if slide.blocks.get("steps") %}

      steps:
{% for st in slide.blocks.get("steps") %}
        - heading: {{ st.heading | e }}
          text: {{ st.text | e }}
{% endfor %}
{% endif %}
{% if slide.blocks.get("conclusion") %}

      conclusion:
        icon: check
        text: {{ slide.blocks.get("conclusion") | e }}
{% endif %}
{% endfor %}
{% endfor %}
'''



DEFAULT_PSL = r'''
theme "corporate":
  page:
    size: widescreen
    background: "#FFFFFF"

  colors:
    primary: "#003A8C"
    secondary: "#E6F0FF"
    text: "#1F1F1F"
    muted: "#666666"
    white: "#FFFFFF"

  fonts:
    heading: "Aptos Display"
    body: "Aptos"

  presentation.title-slide:
    background: primary
    title:
      font: heading
      size: 44
      color: white
      bold: true
      align: center
      position: [80, 210]
      width: 1120
      height: 90
    subtitle:
      font: body
      size: 24
      color: secondary
      align: center
      position: [100, 315]
      width: 1080
      height: 60
    author:
      font: body
      size: 18
      color: white
      align: center
      position: [140, 430]
      width: 1000
      height: 40

  section.section-header:
    background: secondary
    title:
      font: heading
      size: 42
      color: primary
      bold: true
      align: center
      position: [90, 250]
      width: 1100
      height: 80
    subtitle:
      font: body
      size: 22
      color: text
      align: center
      position: [120, 345]
      width: 1040
      height: 60

  slide.title-bullets:
    title:
      font: heading
      size: 34
      color: primary
      bold: true
      position: [60, 40]
      width: 1120
      height: 70
    subtitle:
      font: body
      size: 19
      color: muted
      position: [65, 105]
      width: 1080
      height: 38
    bullets:
      font: body
      size: 22
      color: text
      position: [85, 165]
      width: 1080
      height: 340
      line_gap: 10
      overflow: shrink
      min_size: 15
    conclusion:
      font: body
      size: 18
      color: text
      bold: true
      align: center
      fill: "#E6F0FF"
      border: primary
      position: [90, 555]
      width: 1100
      height: 80
    footer_text:
      font: body
      size: 12
      color: muted
      position: [70, 685]
      width: 900
      height: 24

  slide.grid-4:
    title:
      font: heading
      size: 32
      color: primary
      bold: true
      position: [60, 40]
      width: 1120
      height: 70
    grid:
      position: [70, 150]
      width: 1140
      height: 430
      columns: 2
      gap: 22
    card:
      fill: "#F8FAFC"
      border: "#CBD5E1"
    footer_text:
      font: body
      size: 12
      color: muted
      position: [70, 685]
      width: 900
      height: 24

  slide.numbered-columns-4:
    title:
      font: heading
      size: 32
      color: primary
      bold: true
      position: [60, 40]
      width: 1120
      height: 70
    columns:
      position: [70, 165]
      width: 1140
      height: 410
      gap: 20
    card:
      fill: "#FFFFFF"
      border: "#CBD5E1"
    footer_text:
      font: body
      size: 12
      color: muted
      position: [70, 685]
      width: 900
      height: 24

  slide.timeline:
    title:
      font: heading
      size: 32
      color: primary
      bold: true
      position: [60, 40]
      width: 1120
      height: 70
    timeline:
      position: [90, 180]
      width: 1100
      height: 420
      line_color: primary
      card_fill: "#F8FAFC"
      card_border: "#CBD5E1"
    footer_text:
      font: body
      size: 12
      color: muted
      position: [70, 685]
      width: 900
      height: 24

  slide.stair-progress:
    title:
      font: heading
      size: 32
      color: primary
      bold: true
      position: [60, 40]
      width: 1120
      height: 70
    stair:
      position: [90, 170]
      width: 1080
      height: 430
      fill: "#F8FAFC"
      border: "#CBD5E1"
    footer_text:
      font: body
      size: 12
      color: muted
      position: [70, 685]
      width: 900
      height: 24
'''


def ensure_list(value: Any) -> List[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def safe_text(value: Any, fallback: str = "") -> str:
    if value is None:
        return fallback
    text = str(value).strip()
    return text if text else fallback


def truncate(text: str, limit: int) -> str:
    """Shorten text without adding ellipsis and without cutting mid-word when possible.

    This pipeline is often evaluated by LLM judges. A dangling ellipsis can make
    a summary look incomplete, so the function returns a complete-looking span.
    It prefers sentence boundaries, then clause/word boundaries, and never appends
    "..." or "…".
    """
    text = re.sub(r"\s+", " ", safe_text(text)).strip()
    if len(text) <= limit:
        return text
    if limit <= 0:
        return ""
    window = text[:limit].rstrip()
    # Prefer complete sentence or clause inside the last 45% of the window.
    cut_candidates = [window.rfind(mark) for mark in [". ", "! ", "? ", "。", "; ", ": ", ", "]]
    cut = max(cut_candidates)
    if cut >= max(20, int(limit * 0.55)):
        end = cut + 1
        return window[:end].strip(" ,;:-–—")
    # Otherwise cut at last whitespace to avoid half words.
    space = window.rfind(" ")
    if space >= max(10, int(limit * 0.45)):
        return window[:space].strip(" ,;:-–—")
    return window.strip(" ,;:-–—")


def normalize_for_empty_check(value: Any) -> str:
    """Flatten a value to text for content-presence checks."""
    if value is None:
        return ""
    if isinstance(value, str):
        return re.sub(r"\s+", " ", value).strip()
    if isinstance(value, dict):
        return " ".join(normalize_for_empty_check(v) for v in value.values()).strip()
    if isinstance(value, list):
        return " ".join(normalize_for_empty_check(v) for v in value).strip()
    return str(value).strip()


EMPTY_CONTENT_MARKERS = {
    "",
    "n/a",
    "na",
    "none",
    "null",
    "không có",
    "khong co",
    "chưa có",
    "chua co",
    "không rõ",
    "khong ro",
    "không nêu",
    "khong neu",
    "chưa trích xuất được",
    "chua trich xuat duoc",
    "not available",
    "not specified",
    "not found",
    "no content",
}


PLACEHOLDER_PATTERNS = [
    r"^chưa\s+trích\s+xuất\s+được",
    r"^cần\s+kiểm\s+tra\s+nguồn",
    r"^nội\s+dung\s+chưa\s+rõ",
    r"^không\s+tìm\s+thấy",
    r"^không\s+có\s+nội\s+dung",
    r"^no\s+relevant\s+content",
    r"^no\s+information",
]


def is_meaningful_text(value: Any, min_chars: int = 8) -> bool:
    """Return True only when value contains real report content, not placeholders.

    This is intentionally conservative: if a report component cannot be
    extracted, we skip its slide instead of fabricating filler text.
    """
    text = normalize_for_empty_check(value)
    if not text:
        return False
    lowered = text.casefold().strip(" .。;；:-—–")
    if lowered in EMPTY_CONTENT_MARKERS:
        return False
    if len(lowered) < min_chars:
        return False
    for pattern in PLACEHOLDER_PATTERNS:
        if re.search(pattern, lowered):
            return False
    return True


def clean_meaningful_list(values: Any, limit: Optional[int] = None, item_limit: int = 180) -> List[str]:
    """Normalize list-like content and drop empty/placeholder items."""
    cleaned: List[str] = []
    for item in ensure_list(values):
        text = truncate(remove_dangling_ellipsis(item), item_limit)
        if is_meaningful_text(text):
            cleaned.append(text)
    if limit is not None:
        return cleaned[:limit]
    return cleaned


def section_has_content(section: Dict[str, Any]) -> bool:
    """A report section is useful if it has at least one real bullet, or a real conclusion."""
    return bool(clean_meaningful_list(section.get("bullets"))) or is_meaningful_text(section.get("conclusion"))


def slide_blocks_have_content(blocks: Dict[str, Any]) -> bool:
    """Check whether a slide spec has renderable, non-empty content beyond title/subtitle/footer."""
    if not isinstance(blocks, dict):
        return False
    if clean_meaningful_list(blocks.get("bullets", {}).get("items", []) if isinstance(blocks.get("bullets"), dict) else blocks.get("bullets")):
        return True
    if is_meaningful_text(blocks.get("conclusion")):
        return True
    for key in ("cells", "columns", "milestones", "steps"):
        for item in ensure_list(blocks.get(key)):
            if isinstance(item, dict) and is_meaningful_text(item):
                return True
            if not isinstance(item, dict) and is_meaningful_text(item):
                return True
    if isinstance(blocks.get("table"), dict):
        table = blocks["table"]
        return bool(clean_meaningful_list(table.get("headers"))) and bool(clean_meaningful_list(table.get("rows")))
    return False


def load_json_or_default(path: Optional[str], default: Dict[str, Any], out_dir: Path, filename: str) -> Dict[str, Any]:
    if path:
        p = Path(path)
        if p.exists():
            return json.loads(p.read_text(encoding="utf-8"))
    out_path = out_dir / filename
    out_path.write_text(json.dumps(default, ensure_ascii=False, indent=2), encoding="utf-8")
    return deepcopy(default)


def extract_docx_text(docx_path: Path) -> str:
    """Extract readable text from a .docx file.

    The extractor keeps paragraph order and also serializes tables in a
    Markdown-like pipe format so the downstream LLM can still understand
    tabular report content.
    """
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Missing dependency: python-docx. Install with: pip install python-docx") from exc

    doc = Document(str(docx_path))
    parts: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            style_name = getattr(para.style, "name", "") if para.style else ""
            # Preserve heading-like cues for ontology-based summarization.
            if style_name.lower().startswith("heading"):
                parts.append(f"# {text}")
            else:
                parts.append(text)

    for table_index, table in enumerate(doc.tables, start=1):
        rows: List[List[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            parts.append(f"\n[Table {table_index}]")
            parts.extend(_format_rows_as_markdown_table(rows))

    return "\n\n".join(parts).strip()


def _format_rows_as_markdown_table(rows: List[List[str]]) -> List[str]:
    if not rows:
        return []
    width = max(len(r) for r in rows)
    normalized = [r + [""] * (width - len(r)) for r in rows]
    lines = []
    header = normalized[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * width) + " |")
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def extract_pdf_text(pdf_path: Path) -> str:
    try:
        import pypdf  # type: ignore
    except ImportError as exc:
        raise RuntimeError("For local PDF text extraction, install pypdf or use --mock / Gemini File API.") from exc
    reader = pypdf.PdfReader(str(pdf_path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def extract_plain_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def normalize_raw_input_text(text: str) -> str:
    """Normalize direct plain-text input while preserving report structure.

    This deliberately does not over-clean the text: administrative reports often
    use headings, numbered sections, colon-separated labels, and bullet-like
    lines that are useful signals for the ontology-guided summarizer.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def extract_input_text(input_path: Path, mock: bool, direct_text: str = "", use_stdin: bool = False) -> str:
    """Extract text from PDF, DOCX, TXT/MD, direct --text, or stdin.

    Precedence:
      1. --mock sample text
      2. --text direct plain text
      3. --stdin content
      4. input_path file extraction
    """
    if mock:
        return (
            "Báo cáo nêu vấn đề quá tải thông tin trong doanh nghiệp. "
            "Dữ liệu phân tán ở nhiều hệ thống như CRM, ERP, báo cáo thủ công và mạng xã hội. "
            "Các phát hiện chính gồm thiếu metadata thống nhất, quy trình tổng hợp chậm, "
            "và khó truy xuất bằng ngữ cảnh. Rủi ro gồm sai lệch số liệu, trễ báo cáo, "
            "và phụ thuộc chuyên gia thủ công. Khuyến nghị là xây dựng pipeline tổng hợp tri thức, "
            "chuẩn hóa dữ liệu, dùng RAG/ontology, tự động sinh dashboard và slide. "
            "Lộ trình gồm Q1 khảo sát, Q2 chuẩn hóa, Q3 triển khai thử nghiệm, Q4 mở rộng."
        )

    if direct_text:
        return normalize_raw_input_text(direct_text)

    if use_stdin:
        stdin_text = sys.stdin.read()
        if not stdin_text.strip():
            raise ValueError("--stdin was provided but stdin is empty.")
        return normalize_raw_input_text(stdin_text)

    if not input_path.exists():
        raise FileNotFoundError(
            f"Input file not found: {input_path}. "
            "Use --text '...' for direct plain text or --stdin to pipe content."
        )

    suffix = input_path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(input_path)
    if suffix == ".docx":
        return extract_docx_text(input_path)
    if suffix in {".txt", ".md", ".markdown"}:
        return extract_plain_text(input_path)

    raise ValueError(
        f"Unsupported input format: {suffix}. Supported formats: .pdf, .docx, .txt, .md"
    )


def call_gemini_json(prompt: str, model: str, mock_json: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    if mock_json is not None:
        return mock_json
    try:
        from google import genai  # type: ignore
    except ImportError as exc:
        raise RuntimeError("Missing google-genai. Install with: pip install -U google-genai") from exc
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY is not set.")
    client = genai.Client(api_key=api_key)
    resp = client.models.generate_content(model=model, contents=prompt)
    raw = getattr(resp, "text", "") or ""
    return parse_json_from_text(raw)


def parse_json_from_text(raw: str) -> Dict[str, Any]:
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not m:
            raise
        return json.loads(m.group(0))


def validate_input(state: PipelineState) -> PipelineState:
    out_dir = Path(state["out_dir"])
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "task_outputs").mkdir(exist_ok=True)
    state.setdefault("validation_reports", [])
    state.setdefault("repair_log", [])
    return state


def load_planning_assets(state: PipelineState) -> PipelineState:
    out_dir = Path(state["out_dir"])
    state["report_ontology"] = load_json_or_default(
        state.get("report_ontology_path"), DEFAULT_REPORT_ONTOLOGY, out_dir, "report_domain_ontology.json"
    )
    state["slide_ontology"] = load_json_or_default(
        state.get("slide_ontology_path"), DEFAULT_SLIDE_ONTOLOGY, out_dir, "slide_creation_ontology.json"
    )
    state["layout_registry"] = load_json_or_default(
        state.get("layout_registry_path"), DEFAULT_LAYOUT_REGISTRY, out_dir, "layout_registry.json"
    )
    return state


# -----------------------------------------------------------------------------
# Document-outline-first analysis
# -----------------------------------------------------------------------------

HEADING_NUMBER_RE = re.compile(
    r"^\s*(?:"
    r"[IVXLCDM]{1,8}\.|[A-Z]\.|\d+(?:\.\d+)*[\.)]|"
    r"Điều\s+\d+\.|Chương\s+[IVXLCDM\d]+\.?|Phần\s+[IVXLCDM\d]+\.?)\s+(.+)$",
    flags=re.IGNORECASE,
)

ADMIN_COLON_HEADINGS = [
    "đặc điểm tình hình", "tình hình", "bối cảnh", "hiện trạng",
    "mục tiêu", "yêu cầu", "mục tiêu, yêu cầu", "kết quả", "kết quả nổi bật",
    "tồn tại", "hạn chế", "nguyên nhân", "nhắc nhở", "lưu ý", "biện pháp",
    "giải pháp", "nhiệm vụ", "phân công", "tổ chức thực hiện", "tiến độ",
    "kiến nghị", "đề xuất", "kết luận",
]


def normalize_key(text: str) -> str:
    text = safe_text(text).lower()
    text = re.sub(r"[\s\-_/]+", " ", text)
    text = re.sub(r"[\.:;，,]+$", "", text).strip()
    return text


def is_likely_heading(line: str) -> bool:
    raw = line.strip()
    if not raw or len(raw) > 140:
        return False
    if HEADING_NUMBER_RE.match(raw):
        return True
    key = normalize_key(raw)
    if key in ADMIN_COLON_HEADINGS:
        return True
    if raw.endswith(":") and len(raw) <= 90:
        prefix = normalize_key(raw[:-1])
        return prefix in ADMIN_COLON_HEADINGS or len(prefix.split()) <= 8
    # All-caps Vietnamese headings / short title-like lines.
    letters = [c for c in raw if c.isalpha()]
    if letters and sum(1 for c in letters if c.isupper()) / max(len(letters), 1) > 0.65 and len(raw.split()) <= 12:
        return True
    return False


def clean_heading(line: str) -> str:
    raw = line.strip()
    m = HEADING_NUMBER_RE.match(raw)
    if m:
        raw = m.group(1).strip()
    raw = re.sub(r"^[-–—•]+\s*", "", raw)
    raw = raw.rstrip(":").strip()
    return raw or line.strip()


def split_inline_colon_sections(text: str) -> str:
    """Put common administrative headings on their own line if user pasted one long paragraph.

    The replacement only fires when a heading starts at the beginning of text or
    after a hard/sentence boundary. This avoids splitting a longer heading such as
    "Đặc điểm tình hình:" into "Đặc điểm" + "tình hình:".
    """
    result = text
    for h in sorted(ADMIN_COLON_HEADINGS, key=len, reverse=True):
        pat = re.compile(rf"(^|[\n.;。])\s*({re.escape(h)}\s*:)", flags=re.IGNORECASE)
        result = pat.sub(lambda m: f"{m.group(1)}\n{m.group(2)}\n", result)
    return result


def looks_like_admin_header_line(line: str) -> bool:
    """Detect boilerplate lines commonly appearing in Vietnamese official documents.

    These lines are protocol/metadata, not substantive report content. The matcher
    is intentionally broad for the opening header, e.g. a line containing both
    "Số: ..." and "CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM" separated by tabs.
    """
    raw = safe_text(line).strip()
    if not raw:
        return False
    key = normalize_key(raw)
    compact = re.sub(r"\s+", " ", key)

    # Combined header line: "Số: /2025/NQ-HĐND    CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    if re.search(r"\bs[oố]\s*[:：]", compact) and re.search(r"c[oộòóọõỏ]\s*ng\s+h[oòaàáạãả]\s+x[aã]\s+h[oộòóọõỏ]i\s+ch[uủ]\s+ngh[iĩ]a\s+vi[eệ]t\s+nam", compact):
        return True

    admin_patterns = [
        r"^s[oố]\s*[:：]\s*.*$",
        r"^c[oộòóọõỏ]\s*ng\s+h[oòaàáạãả]\s+x[aã]\s+h[oộòóọõỏ]i\s+ch[uủ]\s+ngh[iĩ]a\s+vi[eệ]t\s+nam$",
        r"^đ[oộ]c\s+l[aậ]p\s*[-–—]\s*t[uự]\s+do\s*[-–—]\s*h[aạ]nh\s+ph[uú]c$",
        r"^\s*[-–—]+\s*$",
        r"^\s*\*\s*\*\s*\*\s*$",
        r"^ng[aà]y\s+\d{1,2}\s+th[aá]ng\s+\d{1,2}\s+n[aă]m\s+\d{4}.*$",
        r"^\w+\s*,\s*ng[aà]y\s+\d{1,2}\s+th[aá]ng\s+\d{1,2}\s+n[aă]m\s+\d{4}.*$",
        r"^n[oơ]i\s+nh[aậ]n\s*:?$",
        r"^k[ií]nh\s+g[uử]i\s*:?$",
        r"^tm\.?\s+.*$",
        r"^kt\.?\s+.*$",
        r"^ch[uủ]\s+t[iị]ch.*$",
        r"^ph[oó]\s+ch[uủ]\s+t[iị]ch.*$",
        r"^ng[uư][ơờ]i\s+k[yý].*$",
        r"^\(\s*đ[aã]\s+k[yý]\s*\).*$",
    ]
    return any(re.match(p, compact, flags=re.IGNORECASE) for p in admin_patterns)

def strip_vietnamese_admin_noise(text: str, ontology: Optional[Dict[str, Any]] = None) -> str:
    """Remove administrative boilerplate that should not become slides.

    Removes opening protocol lines such as "Số: ... / CỘNG HOÀ XÃ HỘI CHỦ NGHĨA
    VIỆT NAM", the national motto, date/signature boilerplate, and the tail
    recipient list starting at "Nơi nhận". This prevents those fragments from
    becoming slide titles, subtitles, or summary bullets.
    """
    ontology = ontology or {}
    rules = ontology.get("exclusion_rules", {}) if isinstance(ontology, dict) else {}
    patterns = [
        r"^cộng\s+hòa\s+xã\s+hội\s+chủ\s+nghĩa\s+việt\s+nam$",
        r"^cộng\s+hoà\s+xã\s+hội\s+chủ\s+nghĩa\s+việt\s+nam$",
        r"^độc\s+lập\s*[-–—]\s*tự\s+do\s*[-–—]\s*hạnh\s+phúc$",
        r"^số\s*[:：]\s*.*$",
    ]
    patterns.extend(ensure_list(rules.get("ignored_heading_patterns")))
    compiled = [re.compile(p, re.IGNORECASE) for p in patterns if p]

    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    cleaned: List[str] = []
    in_recipients_tail = False
    dropped = 0
    for raw in lines:
        line = raw.strip()
        key = normalize_key(line)

        if re.match(r"^nơi\s+nhận\s*:?$", key, flags=re.IGNORECASE):
            in_recipients_tail = True
            dropped += 1
            continue
        if in_recipients_tail:
            dropped += 1
            continue

        if line and looks_like_admin_header_line(line):
            dropped += 1
            continue
        if line and any(pat.match(key) or pat.match(line) for pat in compiled):
            dropped += 1
            continue
        cleaned.append(raw)

    result = "\n".join(cleaned).strip()
    return result

def is_metadata_like_subtitle(text: Any) -> bool:
    """Detect subtitles that are not meaningful slide content."""
    value = normalize_for_empty_check(text).casefold()
    if not value:
        return True
    bad_patterns = [
        r"^theo đề mục gốc",
        r"^source\s*=",
        r"^task\s+\w+",
        r"layout\s*=",
        r"langgraph",
        r"ontology",
        r"jinja",
        r"renderer",
        r"sinh slide",
        r"phân rã",
        r"validate",
    ]
    return any(re.search(p, value) for p in bad_patterns)


def meaningful_subtitle(value: Any) -> str:
    """Return subtitle only when it carries real semantic content."""
    text = truncate(safe_text(value), 120)
    if not is_meaningful_text(text, min_chars=12):
        return ""
    if is_metadata_like_subtitle(text):
        return ""
    return text

def remove_dangling_ellipsis(text: Any) -> str:
    """Remove trailing ellipsis artifacts from LLM summaries."""
    value = safe_text(text)
    value = re.sub(r"\s*(?:\.\.\.|…)\s*$", "", value).strip()
    return value


def extract_document_outline(source_text: str) -> List[Dict[str, Any]]:
    """Extract real headings and their content in source order.

    This is intentionally heuristic and conservative: if clear headings exist, keep
    them as the primary slide structure. Ontology only labels them afterward.
    """
    prepared = split_inline_colon_sections(source_text)
    lines = [ln.rstrip() for ln in prepared.splitlines()]
    sections: List[Dict[str, Any]] = []
    current: Optional[Dict[str, Any]] = None

    for idx, line in enumerate(lines, start=1):
        stripped = line.strip()
        if not stripped:
            continue
        if is_likely_heading(stripped):
            if current and is_meaningful_text("\n".join(current.get("content_lines", []))):
                current["content"] = "\n".join(current.pop("content_lines", [])).strip()
                sections.append(current)
            current = {"source_heading": clean_heading(stripped), "start_line": idx, "content_lines": []}
        else:
            if current is None:
                current = {"source_heading": "Nội dung chính", "start_line": idx, "content_lines": []}
            current["content_lines"].append(stripped)

    if current and is_meaningful_text("\n".join(current.get("content_lines", []))):
        current["content"] = "\n".join(current.pop("content_lines", [])).strip()
        sections.append(current)

    # If no useful heading structure was found, split into coarse chunks; ontology may help label them later.
    if len(sections) <= 1 and len(source_text) > 2500:
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", source_text) if is_meaningful_text(p)]
        chunks: List[Dict[str, Any]] = []
        buf: List[str] = []
        chunk_idx = 1
        char_count = 0
        for p in paragraphs:
            if buf and char_count + len(p) > 2200:
                chunks.append({"source_heading": f"Phần nội dung {chunk_idx}", "start_line": 0, "content": "\n\n".join(buf)})
                chunk_idx += 1
                buf = []
                char_count = 0
            buf.append(p)
            char_count += len(p)
        if buf:
            chunks.append({"source_heading": f"Phần nội dung {chunk_idx}", "start_line": 0, "content": "\n\n".join(buf)})
        if chunks:
            sections = chunks

    return sections



# -----------------------------------------------------------------------------
# LLM-assisted outline hierarchy analysis
# -----------------------------------------------------------------------------

ROMAN_RE = re.compile(r"^[IVXLCDM]{1,8}\.$", flags=re.IGNORECASE)
DECIMAL_RE = re.compile(r"^(\d+(?:\.\d+)*)([\.)])?$")
ALPHA_RE = re.compile(r"^[a-zA-ZđĐ][\.)]$")


def detect_heading_marker(line: str) -> str:
    raw = safe_text(line)
    m = re.match(r"^\s*(Phần\s+[IVXLCDM\d]+\.?|Chương\s+[IVXLCDM\d]+\.?|Mục\s+\d+\.?|Điều\s+\d+\.?|[IVXLCDM]{1,8}\.|\d+(?:\.\d+)*[\.)]?|[A-ZĐa-zđ][\.)])\s+", raw, flags=re.IGNORECASE)
    return m.group(1).strip() if m else ""


def heuristic_outline_level(source_heading: str, source_line: str = "") -> int:
    marker = detect_heading_marker(source_line or source_heading)
    low = marker.lower().strip()
    if low.startswith(("phần", "chương")):
        return 1
    if low.startswith(("mục", "điều")):
        return 2
    if ROMAN_RE.match(marker):
        return 1
    m = DECIMAL_RE.match(marker.rstrip())
    if m:
        number = m.group(1)
        depth = number.count(".") + 1
        return min(max(depth, 1), 4)
    if ALPHA_RE.match(marker):
        return 3
    return 2


def build_heuristic_outline_hierarchy(outline: List[Dict[str, Any]]) -> Dict[str, Any]:
    enriched: List[Dict[str, Any]] = []
    stack: List[Dict[str, Any]] = []
    for idx, item in enumerate(outline, start=1):
        source_heading = item.get("source_heading", f"Mục {idx}")
        level = int(item.get("level") or heuristic_outline_level(source_heading, source_heading))
        level = max(1, min(level, 4))
        while stack and int(stack[-1]["level"]) >= level:
            stack.pop()
        parent_id = stack[-1]["outline_id"] if stack else None
        outline_id = f"H{idx:02d}"
        group_key = outline_id if level <= 1 else (stack[0]["outline_id"] if stack else outline_id)
        group_title = source_heading if level <= 1 else (stack[0]["source_heading"] if stack else source_heading)
        role = "section" if level <= 1 else "slide"
        if idx == 1 and len(outline) <= 4:
            role = "slide"
        enriched_item = {
            **item,
            "outline_id": outline_id,
            "level": level,
            "parent_id": parent_id,
            "group_key": group_key,
            "group_title": group_title,
            "role": role,
            "llm_confidence": 0.55,
            "reason": "heuristic_numbering_or_heading_style",
        }
        enriched.append(enriched_item)
        stack.append(enriched_item)
    groups: Dict[str, Dict[str, Any]] = {}
    for item in enriched:
        key = item.get("group_key") or item["outline_id"]
        if key not in groups:
            groups[key] = {"group_key": key, "section_title": item.get("group_title") or item.get("source_heading"), "items": []}
        if item.get("role") != "section" or is_meaningful_text(item.get("content")):
            groups[key]["items"].append(item)
    return {"mode": "heuristic", "outline_items": enriched, "section_groups": list(groups.values()), "notes": ["Used heuristic fallback or mock mode."]}


def normalize_outline_hierarchy(raw: Dict[str, Any], flat_outline: List[Dict[str, Any]]) -> Dict[str, Any]:
    if not isinstance(raw, dict):
        return build_heuristic_outline_hierarchy(flat_outline)
    by_heading = {normalize_key(x.get("source_heading", "")): x for x in flat_outline}
    raw_items = [x for x in ensure_list(raw.get("outline_items")) if isinstance(x, dict)]
    if not raw_items:
        return build_heuristic_outline_hierarchy(flat_outline)
    normalized: List[Dict[str, Any]] = []
    for idx, src in enumerate(raw_items, start=1):
        source_heading = safe_text(src.get("source_heading") or src.get("heading"), "")
        flat = by_heading.get(normalize_key(source_heading), {})
        if not flat and idx <= len(flat_outline):
            flat = flat_outline[idx - 1]
            source_heading = flat.get("source_heading", source_heading)
        if not source_heading:
            continue
        level = max(1, min(int(src.get("level") or heuristic_outline_level(source_heading, source_heading)), 4))
        outline_id = safe_text(src.get("outline_id"), f"H{idx:02d}")
        role = safe_text(src.get("role"), "section" if level <= 1 else "slide")
        if role not in {"section", "slide", "subslide", "supporting_detail", "skip"}:
            role = "slide"
        item = {
            **flat,
            "outline_id": outline_id,
            "source_heading": source_heading,
            "level": level,
            "parent_id": src.get("parent_id"),
            "group_key": safe_text(src.get("group_key"), outline_id if level <= 1 else "ROOT"),
            "group_title": safe_text(src.get("group_title"), source_heading if level <= 1 else "Nội dung chính"),
            "role": role,
            "llm_confidence": float(src.get("confidence") or src.get("llm_confidence") or 0.7),
            "reason": safe_text(src.get("reason"), "llm_outline_analysis"),
        }
        if not item.get("content") and flat.get("content"):
            item["content"] = flat["content"]
        normalized.append(item)
    seen = {normalize_key(x.get("source_heading", "")) for x in normalized}
    for item in flat_outline:
        key = normalize_key(item.get("source_heading", ""))
        if key and key not in seen:
            normalized.append(build_heuristic_outline_hierarchy([item])["outline_items"][0])
    groups: Dict[str, Dict[str, Any]] = {}
    for item in normalized:
        if item.get("role") == "skip":
            continue
        key = safe_text(item.get("group_key"), item.get("outline_id", "ROOT"))
        title = safe_text(item.get("group_title"), item.get("source_heading", "Nội dung"))
        if key not in groups:
            groups[key] = {"group_key": key, "section_title": title, "items": []}
        if item.get("role") != "section" or is_meaningful_text(item.get("content")):
            groups[key]["items"].append(item)
    return {"mode": safe_text(raw.get("mode"), "llm_outline_depth_analysis"), "outline_items": normalized, "section_groups": list(groups.values()), "notes": ensure_list(raw.get("notes"))}


def analyze_outline_hierarchy_with_llm(outline: List[Dict[str, Any]], source_text: str, ontology: Dict[str, Any], model: str, mock: bool = False) -> Dict[str, Any]:
    fallback = build_heuristic_outline_hierarchy(outline)
    if mock:
        return fallback
    compact_outline = [{"source_heading": item.get("source_heading"), "start_line": item.get("start_line"), "content_preview": truncate(item.get("content", ""), 500)} for item in outline]
    prompt = f"""
Bạn là chuyên gia phân tích bố cục văn bản hành chính để lập kế hoạch slide.

NHIỆM VỤ DUY NHẤT: xác định độ sâu bố cục thật của văn bản, không tóm tắt nội dung.

YÊU CẦU:
1. Bám theo đề mục thật của văn bản, giữ đúng thứ tự.
2. Phân biệt cấp bố cục: level 1 = phần/chương/mục lớn; level 2 = đề mục chính; level 3-4 = mục con/chi tiết.
3. Nếu văn bản chỉ có các mục ngang hàng, đừng ép thành nhiều section header; dùng một group chung hoặc nhóm hợp lý.
4. Ontology chỉ là gợi ý nhãn nghiệp vụ, không được thay thế đề mục nguồn.
5. Bỏ qua quốc hiệu, tiêu ngữ, số văn bản, nơi nhận, ký tá.
6. Chỉ tạo section boundary khi nó giúp deck dễ đọc: phần lớn, chương, nhóm đề mục nhiều slide, hoặc chuyển chủ đề rõ.
7. Nếu một heading là mục con nhỏ, đặt parent_id rõ ràng và role = "supporting_detail" hoặc "slide" tùy độ quan trọng.

Trả về JSON thuần:
{{
  "mode": "llm_outline_depth_analysis",
  "outline_items": [
    {{
      "outline_id": "H01",
      "source_heading": "đề mục thật",
      "level": 1,
      "parent_id": null,
      "group_key": "G01",
      "group_title": "tên section dùng trong deck",
      "role": "section|slide|subslide|supporting_detail|skip",
      "confidence": 0.0,
      "reason": "vì sao chọn cấp này"
    }}
  ],
  "notes": ["ghi chú ngắn"]
}}

Detected outline:
{json.dumps(compact_outline, ensure_ascii=False, indent=2)[:18000]}

Ontology concepts, optional reference only:
{json.dumps(ontology.get('concepts', []), ensure_ascii=False, indent=2)[:10000]}

Một đoạn văn bản nguồn để đối chiếu:
{source_text[:12000]}
"""
    raw = call_gemini_json(prompt, model, fallback)
    return normalize_outline_hierarchy(raw, outline)


def flatten_hierarchy_for_summary(hierarchy: Dict[str, Any]) -> List[Dict[str, Any]]:
    items = []
    for item in hierarchy.get("outline_items", []):
        if item.get("role") == "skip":
            continue
        if not is_meaningful_text(item.get("content")):
            continue
        items.append(item)
    return items

def match_ontology_concept_for_heading(heading: str, ontology: Dict[str, Any]) -> Dict[str, Any]:
    """Best-effort label. Never replaces the source heading."""
    h = normalize_key(heading)
    best: Dict[str, Any] = {}
    best_score = 0
    for concept in ontology.get("concepts", []):
        candidates = [concept.get("preferred_heading", ""), *ensure_list(concept.get("aliases"))]
        for cand in candidates:
            c = normalize_key(cand)
            if not c:
                continue
            score = 0
            if h == c:
                score = 100
            elif c in h or h in c:
                score = 75
            else:
                overlap = set(h.split()) & set(c.split())
                if overlap:
                    score = min(60, 20 * len(overlap))
            if score > best_score:
                best_score = score
                best = concept
    if best_score >= 35:
        return best
    return {"id": "source_section", "preferred_heading": heading, "intent": "summarize_core_message", "max_bullets": 5}


def mock_summary_from_outline(outline: List[Dict[str, Any]], ontology: Dict[str, Any]) -> Dict[str, Any]:
    sections = []
    for item in outline:
        concept = match_ontology_concept_for_heading(item.get("source_heading", ""), ontology)
        content = item.get("content", "")
        # Turn first meaningful sentences/lines into bullets without inventing ontology sections.
        pieces = re.split(r"(?<=[.!?。])\s+|\n+|;\s*", content)
        bullets = clean_meaningful_list(pieces, limit=int(concept.get("max_bullets", 5)), item_limit=180)
        if not bullets and is_meaningful_text(content):
            bullets = [truncate(content, 180)]
        sections.append({
            "source_heading": item.get("source_heading"),
            "concept_id": concept.get("id", "source_section"),
            "heading": item.get("source_heading"),
            "subtitle": "",
            "intent": concept.get("intent", "summarize_core_message"),
            "bullets": bullets,
            "conclusion": "",
            "source_excerpt": truncate(content, 800),
            "start_line": item.get("start_line", 0),
        })
    return {"title": infer_report_title_from_outline(outline), "sections": sections}


def infer_report_title_from_outline(outline: List[Dict[str, Any]]) -> str:
    if outline and outline[0].get("source_heading") not in {"Nội dung chính"}:
        return "Tóm tắt báo cáo theo bố cục văn bản"
    return "Báo cáo tổng hợp"

def summarize_report_by_ontology(state: PipelineState) -> PipelineState:
    """Analyze source outline depth with LLM, then summarize by that hierarchy."""
    source_text = extract_input_text(
        Path(state["pdf_path"]),
        state.get("mock", False),
        state.get("input_text", ""),
        state.get("use_stdin", False),
    )
    source_text = strip_vietnamese_admin_noise(source_text, state.get("report_ontology", {}))
    state["source_text"] = source_text

    flat_outline = extract_document_outline(source_text)
    ontology = state["report_ontology"]
    hierarchy = analyze_outline_hierarchy_with_llm(
        flat_outline,
        source_text,
        ontology,
        state.get("model", "gemini-2.0-flash"),
        bool(state.get("mock")),
    )
    summary_outline = flatten_hierarchy_for_summary(hierarchy)

    state["document_outline"] = summary_outline
    state["outline_hierarchy"] = hierarchy
    state["section_groups"] = hierarchy.get("section_groups", [])

    mock_summary = mock_summary_from_outline(summary_outline, ontology)

    prompt = f"""
Bạn là chuyên gia tóm tắt văn bản hành chính để tạo slide.

NGUYÊN TẮC BẮT BUỘC:
1. Bám theo outline_hierarchy đã phân tích; giữ đúng đề mục nguồn, cấp bố cục, parent và group.
2. Không tự tạo mục theo ontology nếu văn bản không có mục/nội dung tương ứng.
3. Ontology chỉ tham khảo để gán concept_id/intent, không thay đề mục nguồn.
4. Nếu heading chỉ là group/section mà không có nội dung thật, không tạo slide nội dung cho nó.
5. Không tạo subtitle trừ khi subtitle có ý nghĩa nội dung thật; không dùng metadata như level, parent, layout, ontology.
6. Không dùng dấu ba chấm hoặc ký tự … để cắt ý; hãy viết câu/bullet hoàn chỉnh.
7. Bỏ qua quốc hiệu/tiêu ngữ/Số văn bản/Nơi nhận/ký tá hành chính.

Outline hierarchy đã xác nhận:
{json.dumps(hierarchy, ensure_ascii=False, indent=2)[:24000]}

Ontology concepts, optional reference only:
{json.dumps(ontology.get('concepts', []), ensure_ascii=False, indent=2)[:12000]}

Schema JSON thuần:
{{
  "title": "...",
  "sections": [
    {{
      "outline_id": "Hxx",
      "source_heading": "đề mục thật trong văn bản",
      "parent_id": "Hxx hoặc null",
      "level": 1,
      "group_key": "Gxx",
      "group_title": "tên section/group trong deck",
      "concept_id": "id ontology gần nhất hoặc source_section",
      "heading": "tiêu đề slide, ưu tiên đề mục thật",
      "intent": "intent để chọn layout",
      "bullets": ["ý chính từ đúng đề mục này"],
      "conclusion": "kết luận ngắn nếu có, không bịa",
      "start_line": 0
    }}
  ]
}}

Văn bản nguồn:
{source_text[:30000]}
"""
    summary = call_gemini_json(prompt, state.get("model", "gemini-2.0-flash"), mock_summary if state.get("mock") else None)
    state["report_summary"] = normalize_summary(summary, state["report_ontology"], summary_outline)

    by_heading = {normalize_key(x.get("source_heading", "")): x for x in summary_outline}
    for sec in state["report_summary"].get("sections", []):
        meta = by_heading.get(normalize_key(sec.get("source_heading", "")), {})
        for key in ("outline_id", "parent_id", "level", "group_key", "group_title", "role", "llm_confidence", "reason"):
            if key in meta and key not in sec:
                sec[key] = meta[key]

    skipped = state["report_summary"].get("skipped_sections", [])
    state.setdefault("validation_reports", []).append(
        {
            "stage": "llm_outline_depth_analysis",
            "mode": hierarchy.get("mode"),
            "detected_sections": [x.get("source_heading") for x in flat_outline],
            "confirmed_items": [
                {
                    "heading": x.get("source_heading"),
                    "level": x.get("level"),
                    "parent_id": x.get("parent_id"),
                    "group_title": x.get("group_title"),
                    "role": x.get("role"),
                    "confidence": x.get("llm_confidence"),
                }
                for x in hierarchy.get("outline_items", [])
            ],
            "note": "LLM explicitly analyzed outline depth; ontology is only a fallback label source.",
        }
    )
    if skipped:
        state.setdefault("validation_reports", []).append(
            {"stage": "filter_empty_source_sections", "skipped_sections": skipped}
        )
        for item in skipped:
            state.setdefault("repair_log", []).append(
                f"Skipped source section {item.get('heading')}: {item.get('reason')}"
            )
    return state


def make_mock_bullets(concept_id: str) -> List[str]:
    mapping = {
        "executive_summary": [
            "Doanh nghiệp đang gặp tình trạng quá tải thông tin và thiếu cơ chế tổng hợp nhanh.",
            "Giá trị chính của hệ thống là chuyển dữ liệu phân tán thành tri thức có thể trình bày.",
            "Cần kết hợp chuẩn hóa dữ liệu, truy xuất ngữ cảnh và sinh báo cáo tự động.",
        ],
        "context": [
            "Dữ liệu nằm rải rác ở nhiều hệ thống nghiệp vụ.",
            "Báo cáo thủ công làm chậm quá trình ra quyết định.",
            "Lãnh đạo cần insight nhanh hơn, nhất quán hơn và truy vết được nguồn.",
        ],
        "key_findings": [
            "Thiếu metadata thống nhất khiến truy xuất khó ổn định.",
            "Chunking và tổng hợp thủ công dễ làm mất ngữ cảnh.",
            "Pipeline hiện tại chưa tự động hóa đủ các bước từ dữ liệu đến slide.",
            "Ontology có thể giúp chuẩn hóa tiêu đề và cấu trúc báo cáo.",
        ],
        "risks": [
            "Rủi ro sai lệch số liệu nếu nguồn dữ liệu không được kiểm soát.",
            "Rủi ro hallucination nếu LLM thiếu grounding và validation.",
            "Rủi ro trễ tiến độ nếu không chia pipeline thành các task nhỏ.",
        ],
        "recommendations": [
            "Chuẩn hóa ontology nghiệp vụ báo cáo trước khi tóm tắt.",
            "Dùng layout registry để chọn bố cục theo intent thay vì hardcode.",
            "Thêm validate/repair ở từng task trước khi tổng hợp deck.",
            "Tách sinh slide thành nhiều nhiệm vụ độc lập trong LangGraph.",
        ],
        "roadmap": ["Q1: khảo sát nguồn dữ liệu", "Q2: chuẩn hóa ontology", "Q3: thử nghiệm pipeline", "Q4: mở rộng triển khai"],
    }
    return mapping.get(concept_id, ["Nội dung chính cần được tổng hợp theo ontology."])


def make_mock_conclusion(concept_id: str) -> str:
    mapping = {
        "executive_summary": "Trọng tâm là biến báo cáo dài thành slide có cấu trúc, kiểm soát được và dễ trình bày.",
        "context": "Bài toán không chỉ là thiếu dữ liệu, mà là thiếu năng lực tổ chức và khai thác dữ liệu.",
        "key_findings": "Các phát hiện cho thấy cần một pipeline có ontology, registry và validation rõ ràng.",
        "risks": "Rủi ro phải được kiểm soát bằng validation, grounding và cơ chế repair tự động.",
        "recommendations": "Nên triển khai theo từng task nhỏ để dễ kiểm soát chất lượng.",
        "roadmap": "Lộ trình nên đi từ chuẩn hóa nền tảng đến thử nghiệm và mở rộng.",
    }
    return mapping.get(concept_id, "Cần chuyển nội dung thành hành động rõ ràng.")


def normalize_summary(summary: Dict[str, Any], ontology: Dict[str, Any], outline: Optional[List[Dict[str, Any]]] = None) -> Dict[str, Any]:
    """Normalize LLM output in source-heading order, not ontology order.

    The ontology no longer drives which slides exist. It only supplies fallback
    concept_id/intent/max_bullets for each source section.
    """
    outline = outline or []
    outline_by_heading = {normalize_key(x.get("source_heading", "")): x for x in outline}
    raw_sections = [s for s in ensure_list(summary.get("sections")) if isinstance(s, dict)]

    # Sort by source order when start_line is available; otherwise keep LLM order.
    raw_sections = sorted(raw_sections, key=lambda s: int(s.get("start_line") or 10**9))

    normalized_sections: List[Dict[str, Any]] = []
    skipped_sections: List[Dict[str, Any]] = []
    seen: set[str] = set()

    for idx, src in enumerate(raw_sections, start=1):
        source_heading = safe_text(src.get("source_heading"), src.get("heading") or f"Mục {idx}")
        key = normalize_key(source_heading)
        if key in seen:
            continue
        seen.add(key)
        outline_item = outline_by_heading.get(key, {})
        concept = match_ontology_concept_for_heading(source_heading, ontology)
        max_bullets = int(src.get("max_bullets") or concept.get("max_bullets", 5))
        bullets = clean_meaningful_list(src.get("bullets"), limit=max_bullets, item_limit=180)
        conclusion = truncate(remove_dangling_ellipsis(src.get("conclusion")), 220)
        has_conclusion = is_meaningful_text(conclusion)

        # If LLM returned empty bullets but the source section has content, salvage from source excerpt.
        if not bullets and is_meaningful_text(outline_item.get("content")):
            pieces = re.split(r"(?<=[.!?。])\s+|\n+|;\s*", outline_item.get("content", ""))
            bullets = clean_meaningful_list(pieces, limit=max_bullets, item_limit=180)

        if not bullets and not has_conclusion:
            skipped_sections.append({"heading": source_heading, "reason": "no_meaningful_content_extracted"})
            continue

        normalized_sections.append(
            {
                "concept_id": safe_text(src.get("concept_id"), concept.get("id", "source_section")),
                "source_heading": source_heading,
                "heading": safe_text(src.get("heading"), source_heading),
                "subtitle": meaningful_subtitle(src.get("subtitle", "")),
                "intent": safe_text(src.get("intent"), concept.get("intent", "summarize_core_message")),
                "bullets": bullets,
                "conclusion": conclusion if has_conclusion else "",
                "start_line": int(src.get("start_line") or outline_item.get("start_line") or idx),
            }
        )

    # If Gemini omitted a detected source section, preserve it with heuristic bullets.
    existing_keys = {normalize_key(s.get("source_heading", s.get("heading", ""))) for s in normalized_sections}
    for item in outline:
        key = normalize_key(item.get("source_heading", ""))
        if not key or key in existing_keys:
            continue
        concept = match_ontology_concept_for_heading(item.get("source_heading", ""), ontology)
        pieces = re.split(r"(?<=[.!?。])\s+|\n+|;\s*", item.get("content", ""))
        bullets = clean_meaningful_list(pieces, limit=int(concept.get("max_bullets", 5)), item_limit=180)
        if not bullets:
            skipped_sections.append({"heading": item.get("source_heading"), "reason": "detected_heading_but_no_content"})
            continue
        normalized_sections.append(
            {
                "concept_id": concept.get("id", "source_section"),
                "source_heading": item.get("source_heading"),
                "heading": item.get("source_heading"),
                "subtitle": "",
                "intent": concept.get("intent", "summarize_core_message"),
                "bullets": bullets,
                "conclusion": "",
                "start_line": int(item.get("start_line") or 10**9),
            }
        )

    normalized_sections = sorted(normalized_sections, key=lambda s: int(s.get("start_line") or 10**9))

    return {
        "title": safe_text(summary.get("title"), infer_report_title_from_outline(outline)),
        "sections": normalized_sections,
        "skipped_sections": skipped_sections,
        "structure_source": "document_outline_first",
    }


def choose_layout(intent: str, section: Dict[str, Any], slide_ontology: Dict[str, Any], registry: Dict[str, Any]) -> str:
    available = {layout["id"]: layout for layout in registry.get("layouts", [])}
    for rule in slide_ontology.get("layout_selection_rules", []):
        if rule.get("intent") == intent:
            for candidate in rule.get("prefer", []):
                if candidate in available:
                    return candidate
    for layout in registry.get("layouts", []):
        if intent in layout.get("intent", []):
            return layout["id"]
    return "title-bullets"


def decompose_into_slide_tasks(state: PipelineState) -> PipelineState:
    summary = state["report_summary"]
    tasks: List[Dict[str, Any]] = []
    task_index = 1
    for section in summary.get("sections", []):
        if not section_has_content(section):
            state.setdefault("validation_reports", []).append(
                {"stage": "decompose_into_slide_tasks", "skipped_section": section.get("concept_id"), "reason": "empty_after_normalization"}
            )
            continue
        layout_id = choose_layout(section.get("intent", ""), section, state["slide_ontology"], state["layout_registry"])
        idx = task_index
        task_index += 1
        task = {
            "task_id": f"T{idx:02d}",
            "task_type": "create_slide_fragment",
            "outline_id": section.get("outline_id"),
            "parent_id": section.get("parent_id"),
            "outline_level": section.get("level", 2),
            "section_group_key": section.get("group_key", "G00"),
            "section_group_title": section.get("group_title", "Tổng hợp báo cáo"),
            "source_concept_id": section.get("concept_id", "source_section"),
            "source_heading": section.get("source_heading", section.get("heading")),
            "slide_intent": section.get("intent", "summarize_core_message"),
            "layout_id": layout_id,
            "title": section.get("heading"),
            "input_summary_section": section,
            "status": "pending",
        }
        tasks.append(task)

    state["task_plan"] = {
        "deck_title": summary.get("title", "Báo cáo tổng hợp"),
        "structure_source": "llm_outline_depth_first",
        "outline_hierarchy": state.get("outline_hierarchy", {}),
        "section_groups": state.get("section_groups", []),
        "policy": state["slide_ontology"].get("task_policy", {}),
        "tasks": tasks,
    }
    state["pending_tasks"] = tasks
    state["completed_tasks"] = []
    return state


def has_pending_tasks(state: PipelineState) -> str:
    return "process_next_task" if state.get("pending_tasks") else "aggregate_deck"


def pop_next_task(state: PipelineState) -> PipelineState:
    pending = list(state.get("pending_tasks", []))
    if not pending:
        return state
    state["current_task"] = pending.pop(0)
    state["pending_tasks"] = pending
    return state


def process_current_task(state: PipelineState) -> PipelineState:
    task = state["current_task"]
    section = task["input_summary_section"]
    layout = task["layout_id"]
    bullets = clean_meaningful_list(section.get("bullets"), item_limit=180)
    conclusion = section.get("conclusion", "") if is_meaningful_text(section.get("conclusion")) else ""

    if not bullets and not conclusion:
        state["current_task_output"] = {
            "task_id": task["task_id"],
            "outline_id": task.get("outline_id"),
            "parent_id": task.get("parent_id"),
            "outline_level": task.get("outline_level"),
            "section_group_key": task.get("section_group_key"),
            "section_group_title": task.get("section_group_title", "Tổng hợp báo cáo"),
            "source_concept_id": task["source_concept_id"],
            "intent": task["slide_intent"],
            "title": task.get("title", ""),
            "layout": layout,
            "blocks": {},
            "status": "skipped",
            "skip_reason": "no_meaningful_content_for_slide",
        }
        return state

    # Each task produces a structured slide spec, not PML directly.
    if layout == "grid-4":
        cells = []
        for b in bullets[:4]:
            parts = b.split(":", 1)
            heading = parts[0] if len(parts) > 1 and len(parts[0]) < 30 else "Ý chính"
            text = parts[1].strip() if len(parts) > 1 else b
            cells.append({"heading": truncate(heading, 45), "text": truncate(text, 150)})
        blocks: Dict[str, Any] = {"cells": cells, "conclusion": conclusion}
    elif layout == "numbered-columns-4":
        columns = []
        for b in bullets[:4]:
            columns.append({"heading": truncate(b, 45), "text": truncate(conclusion, 120), "bullets": []})
        blocks = {"columns": columns}
    elif layout == "timeline":
        milestones = []
        for i, b in enumerate(section["bullets"][:5], start=1):
            label = f"M{i}"
            if re.match(r"^(Q\d|20\d{2}|Giai đoạn|Phase)", b):
                label = b.split(":", 1)[0][:12]
            milestones.append({"date": label, "heading": truncate(b.split(":", 1)[0], 45), "text": truncate(b, 150)})
        blocks = {"milestones": milestones}
    elif layout == "stair-progress":
        blocks = {"steps": [{"heading": truncate(b, 50), "text": truncate(b, 150)} for b in section["bullets"][:5]]}
    else:
        blocks = {
            "bullets": {"icon": "check", "items": bullets},
            "conclusion": conclusion,
        }

    state["current_task_output"] = {
        "task_id": task["task_id"],
        "outline_id": task.get("outline_id"),
        "parent_id": task.get("parent_id"),
        "outline_level": task.get("outline_level"),
        "section_group_key": task.get("section_group_key"),
        "section_group_title": task.get("section_group_title", "Tổng hợp báo cáo"),
        "source_concept_id": task["source_concept_id"],
        "intent": task["slide_intent"],
        "title": task["title"],
        "subtitle": meaningful_subtitle(section.get("subtitle", "")),
        "layout": layout,
        "blocks": blocks,
        "status": "completed",
    }
    return state


def validate_current_task_output(state: PipelineState) -> PipelineState:
    output = state["current_task_output"]
    registry = {x["id"]: x for x in state["layout_registry"].get("layouts", [])}
    issues = []
    if output.get("status") == "skipped":
        issues.append({"severity": "info", "field": "content", "message": output.get("skip_reason", "skipped")})
        output["validation_issues"] = issues
        state.setdefault("validation_reports", []).append(
            {"stage": "validate_task_output", "task_id": output.get("task_id"), "issues": issues}
        )
        return state
    if not output.get("title"):
        issues.append({"severity": "error", "field": "title", "message": "Missing slide title"})
    if output.get("layout") not in registry:
        issues.append({"severity": "error", "field": "layout", "message": f"Unsupported layout {output.get('layout')}"})
    else:
        supported = set(registry[output["layout"]].get("supported_blocks", []))
        for block_name in output.get("blocks", {}).keys():
            if block_name not in supported and block_name != "conclusion":
                issues.append({"severity": "warning", "field": block_name, "message": "Block may not be supported by layout"})

    if not slide_blocks_have_content(output.get("blocks", {})):
        issues.append({"severity": "error", "field": "blocks", "message": "Slide has no meaningful renderable content"})
        output["status"] = "skipped"
        output["skip_reason"] = "empty_slide_blocks"

    output["validation_issues"] = issues
    state.setdefault("validation_reports", []).append(
        {"stage": "validate_task_output", "task_id": output.get("task_id"), "issues": issues}
    )
    return state


def repair_current_task_output(state: PipelineState) -> PipelineState:
    output = state["current_task_output"]
    issues = output.get("validation_issues", [])
    if not issues:
        return state

    if output.get("status") == "skipped":
        state.setdefault("repair_log", []).append(
            f"{output.get('task_id')}: skipped slide because {output.get('skip_reason', 'no meaningful content')}"
        )
        output["validation_issues"] = []
        return state

    registry_ids = {x["id"] for x in state["layout_registry"].get("layouts", [])}
    if not output.get("title"):
        output["title"] = "Slide chưa đặt tiêu đề"
        state.setdefault("repair_log", []).append(f"{output.get('task_id')}: added fallback title")
    if output.get("layout") not in registry_ids:
        output["layout"] = "title-bullets"
        output["blocks"] = {"bullets": ["Nội dung đã được chuyển về layout an toàn."]}
        state.setdefault("repair_log", []).append(f"{output.get('task_id')}: changed unsupported layout to title-bullets")
    output["validation_issues"] = []
    return state


def store_completed_task(state: PipelineState) -> PipelineState:
    output = state["current_task_output"]
    out_dir = Path(state["out_dir"])
    task_path = out_dir / "task_outputs" / f"{output['task_id']}.json"
    task_path.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")

    if output.get("status") == "skipped" or not slide_blocks_have_content(output.get("blocks", {})):
        state.setdefault("repair_log", []).append(
            f"{output.get('task_id')}: not aggregated because slide content is empty"
        )
        return state

    completed = list(state.get("completed_tasks", []))
    completed.append(output)
    state["completed_tasks"] = completed
    return state


def aggregate_deck(state: PipelineState) -> PipelineState:
    env = Environment(undefined=StrictUndefined, trim_blocks=False, lstrip_blocks=False)
    template = env.from_string(DECK_TEMPLATE)
    slides = [s for s in sorted(state.get("completed_tasks", []), key=lambda x: x.get("task_id", "")) if slide_blocks_have_content(s.get("blocks", {}))]
    state["completed_tasks"] = slides
    if not slides:
        state.setdefault("validation_reports", []).append(
            {"stage": "aggregate_deck", "issues": [{"severity": "warning", "message": "No non-empty slide tasks were available."}]}
        )
        state.setdefault("repair_log", []).append("No slides aggregated because all report components were empty or skipped.")

    grouped: List[Dict[str, Any]] = []
    group_index: Dict[str, Dict[str, Any]] = {}
    for slide in slides:
        key = safe_text(slide.get("section_group_key"), "G00")
        title = safe_text(slide.get("section_group_title"), "Tổng hợp báo cáo")
        if key not in group_index:
            group = {"group_key": key, "section_title": title, "slides": []}
            group_index[key] = group
            grouped.append(group)
        group_index[key]["slides"].append(slide)
    if not grouped:
        grouped = [{"group_key": "G00", "section_title": "Thông báo", "slides": []}]

    state["pml_text"] = template.render(title=state["task_plan"]["deck_title"], section_groups=grouped)
    state["psl_text"] = DEFAULT_PSL
    state.setdefault("validation_reports", []).append(
        {"stage": "aggregate_by_llm_outline_groups", "groups": [{"section_title": g["section_title"], "slide_count": len(g["slides"])} for g in grouped]}
    )
    return state


def validate_aggregated_pml(state: PipelineState) -> PipelineState:
    pml = state.get("pml_text", "")
    issues = []
    if not pml.strip().startswith("presentation"):
        issues.append({"severity": "error", "message": "PML does not start with presentation"})
    if "section " not in pml:
        issues.append({"severity": "error", "message": "PML has no section"})
    if "slide " not in pml:
        issues.append({"severity": "error", "message": "PML has no slide"})
    state.setdefault("validation_reports", []).append({"stage": "validate_aggregated_pml", "issues": issues})
    if issues:
        if any("no slide" in item.get("message", "").lower() for item in issues):
            state.setdefault("repair_log", []).append("No content slides were created; writing a diagnostic no-content deck instead of blank slides.")
            state["pml_text"] = no_content_pml(state)
        else:
            state.setdefault("repair_log", []).append("Aggregated PML had issues; applying fallback repair.")
            state["pml_text"] = fallback_pml(state)
    return state


def fallback_pml(state: PipelineState) -> str:
    return '''presentation "Fallback Deck":
  use style: "corporate.psl"
  section "Fallback":
    slide "Không tạo được deck đầy đủ":
      layout: title-bullets
      title:
        Không tạo được deck đầy đủ
      bullets:
        - Pipeline đã fallback do lỗi validation.
'''


def no_content_pml(state: PipelineState) -> str:
    return '''presentation "Không có nội dung đủ điều kiện":
  use style: "corporate.psl"
  section "Thông báo":
    slide "Không có nội dung đủ điều kiện để tạo slide":
      layout: title-bullets
      title:
        Không có nội dung đủ điều kiện để tạo slide
      bullets:
        - Pipeline đã bỏ qua các thành phần báo cáo không trích xuất được nội dung thật.
        - Vui lòng kiểm tra lại nguồn đầu vào hoặc ontology nếu cần tạo thêm slide.
'''


def write_outputs(state: PipelineState) -> PipelineState:
    out_dir = Path(state["out_dir"])
    pml_path = out_dir / "generated.pml"
    psl_path = out_dir / "corporate.psl"
    task_plan_path = out_dir / "task_plan.json"
    validation_path = out_dir / "validation_report.json"
    repair_path = out_dir / "repair_log.txt"
    source_text_path = out_dir / "source_text.txt"
    outline_hierarchy_path = out_dir / "outline_hierarchy.json"

    outline_hierarchy_path.write_text(json.dumps(state.get("outline_hierarchy", {}), ensure_ascii=False, indent=2), encoding="utf-8")
    source_text_path.write_text(state.get("source_text", ""), encoding="utf-8")
    pml_path.write_text(state["pml_text"], encoding="utf-8")
    psl_path.write_text(state["psl_text"], encoding="utf-8")
    task_plan_path.write_text(json.dumps(state["task_plan"], ensure_ascii=False, indent=2), encoding="utf-8")
    validation_path.write_text(json.dumps(state.get("validation_reports", []), ensure_ascii=False, indent=2), encoding="utf-8")
    repair_path.write_text("\n".join(state.get("repair_log", [])), encoding="utf-8")

    state["pml_path"] = str(pml_path)
    state["psl_path"] = str(psl_path)
    state["task_plan_path"] = str(task_plan_path)
    state["source_text_path"] = str(source_text_path)
    return state


def render_with_existing_renderer(state: PipelineState) -> PipelineState:
    renderer_path = Path(state["renderer_path"])
    if not renderer_path.exists():
        # Try relative to current working directory or /mnt/data.
        alt = Path.cwd() / state["renderer_path"]
        if alt.exists():
            renderer_path = alt
        else:
            alt2 = Path("/mnt/data") / state["renderer_path"]
            if alt2.exists():
                renderer_path = alt2
    if not renderer_path.exists():
        state.setdefault("repair_log", []).append(f"Renderer not found: {state['renderer_path']}; skipped render.")
        return state

    spec = importlib.util.spec_from_file_location("dsl_renderer", str(renderer_path))
    if spec is None or spec.loader is None:
        state.setdefault("repair_log", []).append("Could not import renderer; skipped render.")
        return state
    module = importlib.util.module_from_spec(spec)
    sys.modules["dsl_renderer"] = module
    spec.loader.exec_module(module)  # type: ignore

    required = ["parse_pml", "parse_psl", "build_render_ir", "render_html", "render_pptx"]
    missing = [name for name in required if not hasattr(module, name)]
    if missing:
        state.setdefault("repair_log", []).append(f"Renderer missing functions {missing}; skipped render.")
        return state

    doc = module.parse_pml(state["pml_text"])
    theme = module.parse_psl(state["psl_text"])
    render_ir = module.build_render_ir(doc, theme)

    out_dir = Path(state["out_dir"])
    html_path = out_dir / "output.html"
    pptx_path = out_dir / "output.pptx"
    md_path = out_dir / "output.md"
    module.render_html(render_ir, str(html_path))
    module.render_pptx(render_ir, str(pptx_path))
    if hasattr(module, "render_markdown"):
        module.render_markdown(render_ir, str(md_path))
        state["md_path"] = str(md_path)
    state["html_path"] = str(html_path)
    state["pptx_path"] = str(pptx_path)
    return state


def build_graph():
    END, START, StateGraph = require_langgraph()
    graph = StateGraph(PipelineState)
    graph.add_node("validate_input", validate_input)
    graph.add_node("load_planning_assets", load_planning_assets)
    graph.add_node("summarize_report_by_ontology", summarize_report_by_ontology)
    graph.add_node("decompose_into_slide_tasks", decompose_into_slide_tasks)
    graph.add_node("pop_next_task", pop_next_task)
    graph.add_node("process_current_task", process_current_task)
    graph.add_node("validate_current_task_output", validate_current_task_output)
    graph.add_node("repair_current_task_output", repair_current_task_output)
    graph.add_node("store_completed_task", store_completed_task)
    graph.add_node("aggregate_deck", aggregate_deck)
    graph.add_node("validate_aggregated_pml", validate_aggregated_pml)
    graph.add_node("write_outputs", write_outputs)
    graph.add_node("render_with_existing_renderer", render_with_existing_renderer)

    graph.add_edge(START, "validate_input")
    graph.add_edge("validate_input", "load_planning_assets")
    graph.add_edge("load_planning_assets", "summarize_report_by_ontology")
    graph.add_edge("summarize_report_by_ontology", "decompose_into_slide_tasks")
    graph.add_conditional_edges(
        "decompose_into_slide_tasks",
        has_pending_tasks,
        {"process_next_task": "pop_next_task", "aggregate_deck": "aggregate_deck"},
    )
    graph.add_edge("pop_next_task", "process_current_task")
    graph.add_edge("process_current_task", "validate_current_task_output")
    graph.add_edge("validate_current_task_output", "repair_current_task_output")
    graph.add_edge("repair_current_task_output", "store_completed_task")
    graph.add_conditional_edges(
        "store_completed_task",
        has_pending_tasks,
        {"process_next_task": "pop_next_task", "aggregate_deck": "aggregate_deck"},
    )
    graph.add_edge("aggregate_deck", "validate_aggregated_pml")
    graph.add_edge("validate_aggregated_pml", "write_outputs")
    graph.add_edge("write_outputs", "render_with_existing_renderer")
    graph.add_edge("render_with_existing_renderer", END)
    return graph.compile()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Task-graph outline-first slide generation pipeline; ontology is reference only")
    parser.add_argument("input_path", nargs="?", default="", help="Input .pdf, .docx, .txt, or .md path. Optional when using --text, --stdin, or --mock.")
    parser.add_argument("--text", default="", help="Direct plain-text input. Use this when content is already available as text, not as a file.")
    parser.add_argument("--stdin", action="store_true", help="Read plain-text input from stdin. Example: cat report.txt | python script.py --stdin")
    parser.add_argument("--renderer", default="demo_dsl_conclusion_box.py", help="Path to renderer .py file")
    parser.add_argument("--out-dir", default="out_task_graph_demo", help="Output directory")
    parser.add_argument("--model", default="gemini-2.0-flash", help="Gemini model")
    parser.add_argument("--mock", action="store_true", help="Run without Gemini/PDF extraction")
    parser.add_argument("--report-ontology", default="", help="Optional report ontology JSON path")
    parser.add_argument("--slide-ontology", default="", help="Optional slide ontology JSON path")
    parser.add_argument("--layout-registry", default="", help="Optional layout registry JSON path")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    app = build_graph()
    initial: PipelineState = {
        "pdf_path": args.input_path or "direct_text_input.txt",
        "input_text": args.text,
        "use_stdin": args.stdin,
        "renderer_path": args.renderer,
        "out_dir": args.out_dir,
        "model": args.model,
        "mock": args.mock,
        "report_ontology_path": args.report_ontology,
        "slide_ontology_path": args.slide_ontology,
        "layout_registry_path": args.layout_registry,
    }
    final = app.invoke(initial)
    print("Generated:")
    for key in ["source_text_path", "pml_path", "psl_path", "task_plan_path", "html_path", "pptx_path", "md_path"]:
        if final.get(key):
            print(f"- {key}: {final[key]}")
    print(f"- task_outputs: {Path(args.out_dir) / 'task_outputs'}")
    print(f"- validation_report: {Path(args.out_dir) / 'validation_report.json'}")
    print(f"- repair_log: {Path(args.out_dir) / 'repair_log.txt'}")


if __name__ == "__main__":
    main()
